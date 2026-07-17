import io
import re
from pathlib import Path

from app.schemas.memo import MemoContent

# Kept in sync with frontend/src/lib/legalCopy.ts EXPORT_DISCLAIMER_FOOTER — these
# documents leave the app (shared with investors/donors) with no other app context
# attached, so the caveat has to travel with the file itself.
EXPORT_DISCLAIMER_FOOTER = (
    "This document was generated with AI assistance to support internal diligence and is "
    "provided for informational purposes only. Nitrogen AI is not a registered investment "
    "adviser, broker-dealer, or provider of investment, financial, legal, or tax advice, and "
    "this document does not constitute a recommendation to buy, sell, or hold any investment, "
    "or an offer or solicitation of any kind. Recipients should independently verify all "
    "information and consult qualified professionals before making any investment or business "
    "decision."
)


def _append_disclaimer_footer(doc) -> None:
    """Append the diligence-tool disclaimer as a small italic paragraph at document end."""
    from docx.shared import Pt, RGBColor

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run(EXPORT_DISCLAIMER_FOOTER)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 120, 128)


# Matches inline citation markers like "[1]", "[12]" — kept in sync with the
# citation format the assessment LLM prompts instruct the model to emit
# (see MEMO_SYSTEM_RULES and friends in app/domain/energy/assessments/).
_CITATION_MARKER_RE = re.compile(r"(\[\d+\])")


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """Mark a paragraph as a jump target so citation links can navigate to it.

    Must be called before any paragraph formatting (e.g. left_indent) is set,
    since w:pPr — if later added by python-docx — must remain the first child
    of w:p per the OOXML schema, and bookmarkStart is inserted at index 0 here.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    paragraph._p.insert(0, start)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def _add_internal_hyperlink_run(paragraph, anchor: str, text: str) -> None:
    """Add a run linking to an in-document bookmark (e.g. a References entry)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_external_hyperlink_run(paragraph, url: str, text: str) -> None:
    """Add a run linking to an external URL as a real (clickable) hyperlink."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "005E72")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")  # half-points; 18 == 9pt, matching prior plain-text styling
    r_pr.append(sz)
    run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_text_with_citation_links(paragraph, text: str, citation_numbers: set) -> None:
    """Append text to a paragraph, turning known "[N]" markers into links to References.

    Falls back to a single plain run when there are no citations to link, so
    behavior is unchanged for content without citations.
    """
    if not text:
        return
    if not citation_numbers:
        paragraph.add_run(text)
        return
    for part in _CITATION_MARKER_RE.split(text):
        if not part:
            continue
        match = re.fullmatch(r"\[(\d+)\]", part)
        if match and int(match.group(1)) in citation_numbers:
            _add_internal_hyperlink_run(paragraph, f"cite_{match.group(1)}", part)
        else:
            paragraph.add_run(part)


class DocxExporterService:
    """Service for exporting memos to DOCX format"""
    
    def __init__(self):
        self.template_path = Path(__file__).parent.parent.parent / "templates" / "memo_template.docx"
    
    def generate(
        self,
        memo_content: MemoContent,
        initiative_title: str,
    ) -> bytes:
        """Generate DOCX from memo content (legacy flat structure)"""
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            # Check if we have a template
            if self.template_path.exists():
                logger.info(f"Using template at {self.template_path}")
                return self._generate_from_template(memo_content, initiative_title)
            else:
                logger.info("Template not found, using basic DOCX generation")
                return self._generate_basic(memo_content, initiative_title)
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {str(e)}", exc_info=True)
            raise
    
    def generate_from_sections(
        self,
        memo_content: dict,
        initiative_title: str,
    ) -> bytes:
        """Generate DOCX from memo content with dynamic sections structure"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import logging
        
        logger = logging.getLogger(__name__)
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(memo_content.get("title", "memo"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph(f"Date: {memo_content.get('date', 'N/A')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()
            
            # Recommendation
            recommendation = memo_content.get("recommendation", "hold").upper()
            doc.add_heading("Recommendation", level=1)
            rec_para = doc.add_paragraph()
            rec_run = rec_para.add_run(recommendation)
            rec_run.bold = True
            rec_run.font.size = Pt(14)
            
            doc.add_paragraph()
            
            # Citations — link "[N]" markers in the body to their References entry.
            citations = memo_content.get("citations", [])
            citation_numbers = {
                c.get("number") for c in citations if isinstance(c, dict) and c.get("number") is not None
            }

            # Dynamic sections
            sections = memo_content.get("sections", [])
            for section in sections:
                section_title = section.get("title", "Section")
                section_content = section.get("content", "")
                
                doc.add_heading(section_title, level=1)
                body_para = doc.add_paragraph()
                _add_text_with_citation_links(body_para, section_content, citation_numbers)
                doc.add_paragraph()
            
            if citations:
                doc.add_heading("References", level=1)
                for idx, citation in enumerate(citations):
                    number = citation.get("number", 0)
                    source_type = citation.get("source_type", "evidence")
                    source_title = citation.get("source_title", "Unknown")
                    excerpt = citation.get("excerpt", "")
                    
                    source_label = "[Case Study]" if source_type == "corpus" else "[Evidence]"
                    ref_para = doc.add_paragraph()
                    bookmark_id = number if isinstance(number, int) else 1000 + idx
                    _add_bookmark(ref_para, f"cite_{number}", bookmark_id)
                    ref_para.add_run(f"[{number}] {source_label} {source_title}")
                    
                    if excerpt:
                        excerpt_para = doc.add_paragraph(f'"{excerpt}"')
                        excerpt_para.paragraph_format.left_indent = Inches(0.5)

            _append_disclaimer_footer(doc)

            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output.read()
        
        except Exception as e:
            logger.error(f"Failed to generate DOCX from sections: {str(e)}", exc_info=True)
            raise
    
    def _generate_from_template(
        self,
        memo_content: MemoContent,
        initiative_title: str,
    ) -> bytes:
        """Generate using docxtpl template"""
        from docxtpl import DocxTemplate
        
        doc = DocxTemplate(self.template_path)
        
        # Prepare context
        context = {
            "title": memo_content.title,
            "date": memo_content.date,
            "initiative_title": initiative_title,
            "executive_summary": memo_content.executive_summary,
            "recommendation": memo_content.recommendation.upper(),
            "recommendation_rationale": memo_content.recommendation_rationale,
            "evidence_summary": memo_content.evidence_summary,
            "risks_and_assumptions": memo_content.risks_and_assumptions,
            "open_questions": memo_content.open_questions,
            "citations": [
                {
                    "number": c.number,
                    "source_type": c.source_type,
                    "source_title": c.source_title,
                    "excerpt": c.excerpt,
                }
                for c in memo_content.citations
            ],
        }
        
        doc.render(context)
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
    
    def generate_assessment_docx(
        self,
        content: dict,
        initiative_title: str,
    ) -> bytes:
        """Generate a DOCX for a assessment assessment output (landscape / stakeholder etc.).

        Expects content with keys:
          title, executive_summary, sections[{theme|category, body}],
          strategic_implications|engagement_strategy, recommendations|risk_considerations,
          citations[{number, source_title, source_url, publisher, excerpt}]  (optional)
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import logging

        logger = logging.getLogger(__name__)

        try:
            doc = Document()

            # Citations — collected up front so body paragraphs can turn known
            # "[N]" markers into real hyperlinks that jump to their References entry.
            citations = content.get("citations", [])
            citation_numbers = {
                c.get("number") for c in citations if isinstance(c, dict) and c.get("number") is not None
            }

            # ── Title ──────────────────────────────────────────────────────
            title_heading = doc.add_heading(content.get("title") or initiative_title, 0)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # ── Executive Summary ──────────────────────────────────────────
            exec_sum = content.get("executive_summary", "")
            if exec_sum:
                doc.add_heading("Executive Summary", level=1)
                exec_para = doc.add_paragraph()
                _add_text_with_citation_links(exec_para, exec_sum, citation_numbers)
                doc.add_paragraph()

            # ── Theme / Category sections ──────────────────────────────────
            sections = content.get("sections", [])
            for sec in sections:
                sec_title = sec.get("theme") or sec.get("category") or "Section"
                sec_body = sec.get("body") or sec.get("content") or ""
                if sec_title or sec_body:
                    doc.add_heading(sec_title, level=2)
                    body_para = doc.add_paragraph()
                    _add_text_with_citation_links(body_para, sec_body, citation_numbers)
                    doc.add_paragraph()

            # ── Trailing sections (any scalar extra keys) ──────────────────
            trailing_keys = [
                ("strategic_implications", "Strategic Implications"),
                ("engagement_strategy", "Engagement Strategy"),
                ("engagement_recommendations", "Engagement Recommendations"),
                ("recommendations", "Recommendations & Next Steps"),
                ("risk_considerations", "Risk Considerations"),
            ]
            for key, label in trailing_keys:
                val = content.get(key)
                if val:
                    doc.add_heading(label, level=1)
                    if isinstance(val, list):
                        for item in val:
                            item_para = doc.add_paragraph()
                            item_para.add_run("• ")
                            _add_text_with_citation_links(item_para, str(item), citation_numbers)
                    else:
                        val_para = doc.add_paragraph()
                        _add_text_with_citation_links(val_para, val, citation_numbers)
                    doc.add_paragraph()

            # ── References ─────────────────────────────────────────────────
            if citations:
                doc.add_heading("References", level=1)
                for idx, cit in enumerate(citations):
                    num = cit.get("number", "")
                    source_title = cit.get("source_title", "Unknown source")
                    source_url = cit.get("source_url") or ""
                    publisher = cit.get("publisher") or ""
                    excerpt = cit.get("excerpt") or ""

                    ref_para = doc.add_paragraph()
                    bookmark_id = num if isinstance(num, int) else 1000 + idx
                    _add_bookmark(ref_para, f"cite_{num}", bookmark_id)
                    line = f"[{num}] {source_title}"
                    if publisher:
                        line += f" — {publisher}"
                    ref_para.add_run(line)
                    if source_url:
                        url_para = doc.add_paragraph()
                        _add_external_hyperlink_run(url_para, source_url, source_url)
                        url_para.paragraph_format.left_indent = Inches(0.3)
                    if excerpt:
                        excerpt_para = doc.add_paragraph(f'"{excerpt[:250]}…"')
                        excerpt_para.paragraph_format.left_indent = Inches(0.3)
                        for run in excerpt_para.runs:
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(90, 90, 96)

            _append_disclaimer_footer(doc)

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output.read()

        except Exception as e:
            logger.error(f"Failed to generate assessment DOCX: {e}", exc_info=True)
            raise

    def _generate_basic(
        self,
        memo_content: MemoContent,
        initiative_title: str,
    ) -> bytes:
        """Generate basic DOCX without template"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading(memo_content.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Date: {memo_content.date}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()

        citation_numbers = {c.number for c in memo_content.citations if c.number is not None}

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        exec_para = doc.add_paragraph()
        _add_text_with_citation_links(exec_para, memo_content.executive_summary, citation_numbers)
        
        # Recommendation
        doc.add_heading("Recommendation", level=1)
        rec_para = doc.add_paragraph()
        rec_run = rec_para.add_run(memo_content.recommendation.upper())
        rec_run.bold = True
        rec_run.font.size = Pt(14)
        
        rationale_para = doc.add_paragraph()
        _add_text_with_citation_links(rationale_para, memo_content.recommendation_rationale, citation_numbers)
        
        # Evidence Summary
        doc.add_heading("Evidence Summary", level=1)
        evidence_para = doc.add_paragraph()
        _add_text_with_citation_links(evidence_para, memo_content.evidence_summary, citation_numbers)
        
        # Risks and Variables
        doc.add_heading("Risks and Variables", level=1)
        risks_para = doc.add_paragraph()
        _add_text_with_citation_links(risks_para, memo_content.risks_and_assumptions, citation_numbers)
        
        # Open Questions
        if memo_content.open_questions:
            doc.add_heading("Open Questions", level=1)
            for question in memo_content.open_questions:
                q_para = doc.add_paragraph()
                q_para.add_run("• ")
                _add_text_with_citation_links(q_para, question, citation_numbers)
        
        # Citations
        if memo_content.citations:
            doc.add_heading("References", level=1)
            for idx, citation in enumerate(memo_content.citations):
                source_label = "[Case Study]" if citation.source_type == "corpus" else "[Evidence]"
                ref_para = doc.add_paragraph()
                bookmark_id = citation.number if isinstance(citation.number, int) else 1000 + idx
                _add_bookmark(ref_para, f"cite_{citation.number}", bookmark_id)
                ref_para.add_run(f"[{citation.number}] {source_label} {citation.source_title}")
                excerpt_para = doc.add_paragraph(f'"{citation.excerpt}"')
                excerpt_para.paragraph_format.left_indent = Inches(0.5)

        _append_disclaimer_footer(doc)

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
