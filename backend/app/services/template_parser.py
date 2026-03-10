"""Parse DOCX and XLSX templates into a structured representation of their
requirements — sections, fields, placeholders, formulas, etc."""

from __future__ import annotations

import io
import re
import uuid
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# Patterns that suggest a placeholder/blank in a DOCX template
_PLACEHOLDER_RE = re.compile(
    r"\[_{2,}\]"       # [___]
    r"|{{[^}]+}}"      # {{field_name}}
    r"|\<[^>]+\>"      # <field_name>
    r"|\[TBD\]"        # [TBD]
    r"|\[INSERT\b[^\]]*\]"  # [INSERT ...]
    r"|\[ENTER\b[^\]]*\]"   # [ENTER ...]
    , re.IGNORECASE,
)


@dataclass
class TemplateField:
    id: str
    label: str
    description: str
    field_type: str  # text | number | narrative | table_row | formula
    location: str    # e.g. "para:3", "table:0:row:2:col:1", "Sheet1!B5"
    required: bool = True
    is_calculated: bool = False

    def to_dict(self) -> dict:
        return {
            "id": self.id,
            "label": self.label,
            "description": self.description,
            "field_type": self.field_type,
            "location": self.location,
            "required": self.required,
            "is_calculated": self.is_calculated,
        }


@dataclass
class TemplateSection:
    id: str
    title: str
    fields: list[TemplateField] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "id": self.id,
            "title": self.title,
            "fields": [f.to_dict() for f in self.fields],
        }


@dataclass
class TemplateStructure:
    file_type: str   # "docx" | "xlsx"
    sections: list[TemplateSection] = field(default_factory=list)
    raw_text: str = ""

    def to_dict(self) -> dict:
        return {
            "file_type": self.file_type,
            "sections": [s.to_dict() for s in self.sections],
            "raw_text": self.raw_text[:8000],
        }

    @property
    def total_fields(self) -> int:
        return sum(len(s.fields) for s in self.sections)


class TemplateParserService:
    """Parse DOCX / XLSX bytes into a TemplateStructure."""

    # ── DOCX ────────────────────────────────────────────────────────

    def parse_docx_template(self, content: bytes) -> TemplateStructure:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = Document(io.BytesIO(content))
        sections: list[TemplateSection] = []
        current_section: TemplateSection | None = None
        raw_parts: list[str] = []

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            raw_parts.append(text)

            is_heading = para.style and para.style.name and para.style.name.startswith("Heading")

            if is_heading:
                current_section = TemplateSection(
                    id=str(uuid.uuid4())[:8],
                    title=text,
                )
                sections.append(current_section)
                continue

            if current_section is None:
                current_section = TemplateSection(id="intro", title="Introduction")
                sections.append(current_section)

            placeholders = _PLACEHOLDER_RE.findall(text)
            if placeholders:
                for ph in placeholders:
                    current_section.fields.append(TemplateField(
                        id=str(uuid.uuid4())[:8],
                        label=ph.strip("[]<>{}_ ") or f"Field in paragraph {idx}",
                        description=text[:120],
                        field_type="text",
                        location=f"para:{idx}",
                    ))

        for t_idx, table in enumerate(doc.tables):
            table_section = TemplateSection(
                id=f"table_{t_idx}",
                title=f"Table {t_idx + 1}",
            )
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        label_text = ""
                        if c_idx > 0:
                            label_text = row.cells[0].text.strip()
                        if not label_text and r_idx > 0:
                            label_text = table.rows[0].cells[c_idx].text.strip()
                        if label_text:
                            table_section.fields.append(TemplateField(
                                id=str(uuid.uuid4())[:8],
                                label=label_text,
                                description=f"Empty cell at row {r_idx + 1}, column {c_idx + 1}",
                                field_type="text",
                                location=f"table:{t_idx}:row:{r_idx}:col:{c_idx}",
                            ))
                    elif _PLACEHOLDER_RE.search(cell_text):
                        table_section.fields.append(TemplateField(
                            id=str(uuid.uuid4())[:8],
                            label=cell_text[:80],
                            description=f"Placeholder in table {t_idx + 1}",
                            field_type="text",
                            location=f"table:{t_idx}:row:{r_idx}:col:{c_idx}",
                        ))
                    raw_parts.append(cell_text)

            if table_section.fields:
                sections.append(table_section)

        if not sections:
            sections.append(TemplateSection(id="full_doc", title="Document"))

        return TemplateStructure(
            file_type="docx",
            sections=sections,
            raw_text="\n".join(raw_parts),
        )

    # ── XLSX ────────────────────────────────────────────────────────

    def parse_xlsx_template(self, content: bytes) -> TemplateStructure:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(content), data_only=False)
        sections: list[TemplateSection] = []
        raw_parts: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            section = TemplateSection(id=sheet_name, title=sheet_name)

            header_row: dict[int, str] = {}
            label_col: dict[int, str] = {}

            for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=ws.max_row), start=1):
                for cell in row:
                    if cell.value is not None and isinstance(cell.value, str):
                        val = cell.value.strip()
                        if val:
                            if row_idx == 1:
                                header_row[cell.column] = val
                            if cell.column == 1:
                                label_col[row_idx] = val
                            raw_parts.append(val)

            for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=ws.max_row), start=1):
                for cell in row:
                    col_letter = cell.column_letter
                    loc = f"{sheet_name}!{col_letter}{row_idx}"

                    if cell.data_type == "f" or (
                        isinstance(cell.value, str) and cell.value.startswith("=")
                    ):
                        label = header_row.get(cell.column, "") or label_col.get(row_idx, "")
                        if label:
                            section.fields.append(TemplateField(
                                id=str(uuid.uuid4())[:8],
                                label=label,
                                description=f"Formula at {loc}: {cell.value}",
                                field_type="formula",
                                location=loc,
                                required=False,
                                is_calculated=True,
                            ))
                        continue

                    if cell.value is None and row_idx > 1 and cell.column > 1:
                        label = header_row.get(cell.column, "") or label_col.get(row_idx, "")
                        if label:
                            section.fields.append(TemplateField(
                                id=str(uuid.uuid4())[:8],
                                label=label,
                                description=f"Input cell at {loc}",
                                field_type="number" if _looks_numeric(label) else "text",
                                location=loc,
                            ))

            if section.fields:
                sections.append(section)

        if not sections:
            sections.append(TemplateSection(id="sheet1", title="Sheet1"))

        return TemplateStructure(
            file_type="xlsx",
            sections=sections,
            raw_text="\n".join(raw_parts),
        )


def _looks_numeric(label: str) -> bool:
    keywords = [
        "amount", "cost", "price", "rate", "total", "qty", "quantity",
        "capex", "opex", "revenue", "budget", "kw", "mw", "kwh", "mwh",
        "capacity", "factor", "ratio", "percent", "%", "years", "months",
    ]
    lower = label.lower()
    return any(k in lower for k in keywords)
