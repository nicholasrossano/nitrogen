"""Fill DOCX and XLSX templates with resolved requirement values while
preserving all original formatting."""

from __future__ import annotations

import io
import re
import logging
from dataclasses import dataclass

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(
    r"\[_{2,}\]"
    r"|{{([^}]+)}}"
    r"|\<([^>]+)\>"
    r"|\[TBD\]"
    r"|\[INSERT\b[^\]]*\]"
    r"|\[ENTER\b[^\]]*\]"
    , re.IGNORECASE,
)


@dataclass
class FillValue:
    location: str
    value: str


class TemplateFillerService:
    """Populate a DOCX or XLSX template in-place, preserving formatting."""

    # ── DOCX ────────────────────────────────────────────────────────

    def fill_docx(
        self,
        template_bytes: bytes,
        requirements: list[dict],
    ) -> bytes:
        """Fill a DOCX template and return the modified bytes.

        `requirements` is the list of RequirementStatus.to_dict() dicts.
        """
        from docx import Document

        doc = Document(io.BytesIO(template_bytes))

        value_by_location: dict[str, str] = {}
        for req in requirements:
            if req.get("value") and req.get("source_location"):
                value_by_location[req["source_location"]] = req["value"]

        label_to_value: dict[str, str] = {}
        for req in requirements:
            if req.get("value") and req.get("label"):
                label_to_value[req["label"].lower().strip()] = req["value"]

        for idx, para in enumerate(doc.paragraphs):
            loc = f"para:{idx}"
            if loc in value_by_location:
                self._replace_placeholders_in_paragraph(para, value_by_location[loc])
            elif _PLACEHOLDER_RE.search(para.text):
                self._smart_replace_paragraph(para, label_to_value)

        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    loc = f"table:{t_idx}:row:{r_idx}:col:{c_idx}"
                    if loc in value_by_location:
                        for p in cell.paragraphs:
                            self._set_paragraph_text(p, value_by_location[loc])
                        continue
                    for p in cell.paragraphs:
                        if _PLACEHOLDER_RE.search(p.text):
                            self._smart_replace_paragraph(p, label_to_value)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _replace_placeholders_in_paragraph(para, value: str):
        """Replace all placeholder patterns in a paragraph with a value,
        preserving run-level formatting."""
        full_text = para.text
        new_text = _PLACEHOLDER_RE.sub(value, full_text)
        if new_text != full_text:
            TemplateFillerService._set_paragraph_text(para, new_text)

    @staticmethod
    def _smart_replace_paragraph(para, label_to_value: dict[str, str]):
        """Try to match a placeholder to a label-keyed value."""
        full_text = para.text
        def _replacer(m: re.Match) -> str:
            matched = m.group(0)
            key = matched.strip("[]<>{}_ ").lower()
            return label_to_value.get(key, matched)
        new_text = _PLACEHOLDER_RE.sub(_replacer, full_text)
        if new_text != full_text:
            TemplateFillerService._set_paragraph_text(para, new_text)

    @staticmethod
    def _set_paragraph_text(para, text: str):
        """Set paragraph text while preserving formatting of the first run."""
        if not para.runs:
            para.text = text
            return
        first_run = para.runs[0]
        for run in para.runs[1:]:
            run.text = ""
        first_run.text = text

    # ── XLSX ────────────────────────────────────────────────────────

    def fill_xlsx(
        self,
        template_bytes: bytes,
        requirements: list[dict],
    ) -> bytes:
        """Fill an XLSX template and return the modified bytes.

        Formula cells are left untouched — they recalculate when the
        spreadsheet is opened in Excel/Sheets.
        """
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(template_bytes))

        value_by_location: dict[str, str] = {}
        for req in requirements:
            if req.get("value") and req.get("source_location") and not req.get("is_calculated"):
                value_by_location[req["source_location"]] = req["value"]

        for loc, val in value_by_location.items():
            try:
                if "!" not in loc:
                    continue
                sheet_name, cell_ref = loc.split("!", 1)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    numeric = self._try_numeric(val)
                    ws[cell_ref] = numeric if numeric is not None else val
            except Exception:
                logger.warning("Failed to fill cell %s", loc, exc_info=True)

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    @staticmethod
    def _try_numeric(value: str) -> float | int | None:
        """Attempt to parse a string as a number."""
        cleaned = value.replace(",", "").replace(" ", "").strip()
        try:
            if "." in cleaned:
                return float(cleaned)
            return int(cleaned)
        except (ValueError, TypeError):
            return None
