"""
Gold Standard Cover Letter Service

Fills the original DOCX template with user-provided field values,
computes completion status, and produces export-ready documents.
"""

import io
import logging
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _get_fallback_field_schema() -> list[dict]:
    """Return the hardcoded GS Cover Letter field schema as plain dicts.

    Used as a fast fallback when the template cannot be fetched from the web
    and no cached version exists in the database.
    """
    loc = lambda i: {"type": "paragraph", "index": i}
    fields = [
        {"field_id": "project_title", "label": "Project Title", "field_type": "text",
         "section": "Project Information", "required": True, "docx_location": loc(0),
         "placeholder_text": "[Enter Project Title]", "help_text": ""},
        {"field_id": "gs_id", "label": "Gold Standard ID", "field_type": "text",
         "section": "Project Information", "required": False, "docx_location": loc(1),
         "placeholder_text": "[GS ID if assigned]", "help_text": "Leave blank if not yet assigned"},
        {"field_id": "project_developer_name", "label": "Project Developer Name", "field_type": "text",
         "section": "Project Developer", "required": True, "docx_location": loc(2),
         "placeholder_text": "[Developer Name]", "help_text": ""},
        {"field_id": "project_developer_address", "label": "Address", "field_type": "multiline",
         "section": "Project Developer", "required": True, "docx_location": loc(3),
         "placeholder_text": "[Address]", "help_text": ""},
        {"field_id": "project_developer_contact", "label": "Contact Person", "field_type": "text",
         "section": "Project Developer", "required": True, "docx_location": loc(4),
         "placeholder_text": "[Contact Name]", "help_text": ""},
        {"field_id": "project_developer_email", "label": "Email", "field_type": "text",
         "section": "Project Developer", "required": True, "docx_location": loc(5),
         "placeholder_text": "[email@example.com]", "help_text": ""},
        {"field_id": "project_developer_phone", "label": "Phone", "field_type": "text",
         "section": "Project Developer", "required": False, "docx_location": loc(6),
         "placeholder_text": "[+1 ...]", "help_text": ""},
        {"field_id": "methodology", "label": "Methodology / Protocol", "field_type": "text",
         "section": "Project Details", "required": True, "docx_location": loc(7),
         "placeholder_text": "[Methodology name]", "help_text": ""},
        {"field_id": "project_country", "label": "Host Country", "field_type": "text",
         "section": "Project Details", "required": True, "docx_location": loc(8),
         "placeholder_text": "[Country]", "help_text": ""},
        {"field_id": "project_scale", "label": "Project Scale", "field_type": "choice",
         "section": "Project Details", "required": True, "docx_location": loc(9),
         "placeholder_text": "[Micro / Small / Regular / Large]",
         "help_text": "Micro, Small, Regular, or Large"},
        {"field_id": "crediting_period_start", "label": "Crediting Period Start Date",
         "field_type": "date", "section": "Project Details", "required": True,
         "docx_location": loc(10), "placeholder_text": "[YYYY-MM-DD]", "help_text": ""},
        {"field_id": "crediting_period_end", "label": "Crediting Period End Date",
         "field_type": "date", "section": "Project Details", "required": False,
         "docx_location": loc(11), "placeholder_text": "[YYYY-MM-DD]", "help_text": ""},
        {"field_id": "estimated_annual_credits",
         "label": "Estimated Annual Emission Reductions (tCO2e)",
         "field_type": "text", "section": "Project Details", "required": False,
         "docx_location": loc(12), "placeholder_text": "[e.g. 10,000]", "help_text": ""},
        {"field_id": "project_description_summary", "label": "Brief Project Description",
         "field_type": "multiline", "section": "Project Summary", "required": True,
         "docx_location": loc(13), "placeholder_text": "[Brief description]", "help_text": ""},
        {"field_id": "sustainability_contributions",
         "label": "Sustainability Development Contributions",
         "field_type": "multiline", "section": "Sustainability", "required": False,
         "docx_location": loc(14), "placeholder_text": "[SDG contributions]", "help_text": ""},
        {"field_id": "stakeholder_engagement_summary",
         "label": "Stakeholder Engagement Summary",
         "field_type": "multiline", "section": "Stakeholder Engagement", "required": False,
         "docx_location": loc(15), "placeholder_text": "[Engagement summary]", "help_text": ""},
        {"field_id": "documents_submitted",
         "label": "Documents Submitted with this Cover Letter",
         "field_type": "multiline", "section": "Submission Details", "required": True,
         "docx_location": loc(16), "placeholder_text": "[List documents]", "help_text": ""},
        {"field_id": "signatory_name", "label": "Authorized Signatory Name",
         "field_type": "text", "section": "Declaration", "required": True,
         "docx_location": loc(17), "placeholder_text": "[Full name]", "help_text": ""},
        {"field_id": "signatory_title", "label": "Title / Position",
         "field_type": "text", "section": "Declaration", "required": True,
         "docx_location": loc(18), "placeholder_text": "[Title]", "help_text": ""},
        {"field_id": "signatory_organization", "label": "Organization",
         "field_type": "text", "section": "Declaration", "required": True,
         "docx_location": loc(19), "placeholder_text": "[Organization]", "help_text": ""},
        {"field_id": "signature_date", "label": "Date",
         "field_type": "date", "section": "Declaration", "required": True,
         "docx_location": loc(20), "placeholder_text": "[YYYY-MM-DD]", "help_text": ""},
        {"field_id": "signature", "label": "Signature",
         "field_type": "signature", "section": "Declaration", "required": True,
         "docx_location": loc(21), "placeholder_text": "[Signature]",
         "help_text": "Leave blank — sign after export"},
    ]
    return fields


@dataclass
class FieldCompletion:
    field_id: str
    label: str
    required: bool
    status: str  # empty, filled, signature_pending


@dataclass
class CompletionStatus:
    total_fields: int
    filled_fields: int
    required_fields: int
    required_filled: int
    status: str  # not_started, in_progress, complete, ready_for_signature
    fields: list[FieldCompletion]


GS_CHECKLIST_ITEMS = [
    {
        "id": "cover_letter",
        "name": "Cover Letter",
        "supported": True,
        "description": "Official project cover letter with developer details, project summary, and signatory declaration. Required for all GS4GG submissions.",
        "template_url": None,
    },
    {
        "id": "preliminary_review",
        "name": "Preliminary Review Submission",
        "supported": True,
        "description": "Initial project screening submission for Gold Standard review prior to full design certification.",
        "template_url": "https://globalgoals.goldstandard.org/standards/",
    },
    {
        "id": "pdd",
        "name": "Project Design Document (PDD)",
        "supported": False,
        "description": "Comprehensive project design document describing methodology, baseline, monitoring plan, and expected emission reductions.",
        "template_url": "https://globalgoals.goldstandard.org/standards/",
    },
    {
        "id": "stakeholder_report",
        "name": "Stakeholder Consultation Report",
        "supported": False,
        "description": "Report on local stakeholder consultations conducted as part of project design, including feedback and responses.",
        "template_url": "https://globalgoals.goldstandard.org/standards/",
    },
    {
        "id": "sdg_impact_tool",
        "name": "SDG Impact Tool",
        "supported": False,
        "description": "Digital tool for assessing project contributions to UN Sustainable Development Goals. Completed online via the GS SDG tool.",
        "template_url": "https://sdgimpact.goldstandard.org/",
    },
    {
        "id": "oda_declaration",
        "name": "ODA Declaration Form",
        "supported": False,
        "conditional": True,
        "description": "Declaration that the project does not receive Official Development Assistance (ODA) funding. Required when the project involves public funding.",
        "template_url": "https://globalgoals.goldstandard.org/standards/",
    },
    {
        "id": "terms_conditions",
        "name": "Terms & Conditions Acknowledgement",
        "supported": False,
        "description": "Acknowledgement of Gold Standard platform terms and conditions for project registration.",
        "template_url": "https://globalgoals.goldstandard.org/",
    },
]


class CoverLetterService:
    def fill_template(self, docx_bytes: bytes, field_schema: list[dict], field_values: dict) -> bytes:
        """Fill the original DOCX template with user values, preserving formatting.

        Uses each field's ``docx_location`` to navigate directly to the correct
        paragraph, table cell, or content control — avoiding collisions when the
        same placeholder text (e.g. "[Inert here]") appears in many locations.
        """
        doc = DocxDocument(io.BytesIO(docx_bytes))

        for fdef in field_schema:
            fid = fdef["field_id"]
            entry = field_values.get(fid)
            if not entry or not entry.get("value"):
                continue

            value = entry["value"]
            placeholder = fdef.get("placeholder_text", "")
            if not placeholder:
                continue

            loc = fdef.get("docx_location", {})
            loc_type = loc.get("type")

            if loc_type == "table_cell":
                try:
                    table = doc.tables[loc["table"]]
                    cell = table.rows[loc["row"]].cells[loc["col"]]
                    for para in cell.paragraphs:
                        if placeholder in para.text:
                            self._replace_in_runs(para, placeholder, value)
                except (IndexError, KeyError):
                    logger.warning("Could not locate DOCX table cell for field %s", fid)

            elif loc_type == "paragraph":
                for para in doc.paragraphs:
                    if placeholder in para.text:
                        self._replace_in_runs(para, placeholder, value)
                        break

            elif loc_type == "content_control":
                tag = loc.get("tag")
                if tag:
                    for sdt in doc.element.body.iter(qn('w:sdt')):
                        tag_el = sdt.find(f'.//{qn("w:tag")}')
                        if tag_el is not None and tag_el.get(qn('w:val')) == tag:
                            for t_el in sdt.iter(qn('w:t')):
                                t_el.text = value
                                break
                            break

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def get_completion_status(self, field_schema: list[dict], field_values: dict) -> CompletionStatus:
        """Compute per-field and overall completion status."""
        fields: list[FieldCompletion] = []
        filled = 0
        required_count = 0
        required_filled = 0
        has_signature_field = False

        for fdef in field_schema:
            fid = fdef["field_id"]
            entry = field_values.get(fid, {})
            value = entry.get("value") if isinstance(entry, dict) else None
            is_required = fdef.get("required", False)
            is_signature = fdef.get("field_type") == "signature"

            if is_signature:
                has_signature_field = True
                status = "signature_pending"
                if value:
                    status = "filled"
                    filled += 1
                    if is_required:
                        required_filled += 1
            elif value:
                status = "filled"
                filled += 1
                if is_required:
                    required_filled += 1
            else:
                status = "empty"

            if is_required:
                required_count += 1

            fields.append(FieldCompletion(
                field_id=fid,
                label=fdef.get("label", fid),
                required=is_required,
                status=status,
            ))

        total = len(field_schema)
        if filled == 0:
            overall = "not_started"
        elif required_filled >= required_count and has_signature_field:
            sig_fields = [f for f in fields if f.status == "signature_pending"]
            overall = "ready_for_signature" if sig_fields else "complete"
        elif required_filled >= required_count:
            overall = "complete"
        else:
            overall = "in_progress"

        return CompletionStatus(
            total_fields=total,
            filled_fields=filled,
            required_fields=required_count,
            required_filled=required_filled,
            status=overall,
            fields=fields,
        )

    @staticmethod
    def _replace_in_runs(paragraph, old_text: str, new_text: str):
        """Replace text across runs while preserving the formatting of the first run."""
        runs = paragraph.runs
        if not runs:
            return

        combined = ''.join(r.text for r in runs)
        if old_text not in combined:
            return

        new_combined = combined.replace(old_text, new_text, 1)

        if runs:
            runs[0].text = new_combined
            for r in runs[1:]:
                r.text = ""
