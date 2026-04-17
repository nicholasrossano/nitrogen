"""Decision log builders for module-level and initiative-level reporting."""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any
from uuid import UUID

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.decision_event import DecisionEvent
from app.models.initiative import Initiative
from app.models.module_instance import ModuleInstance
from app.modules.base import DecisionLogAttribution
from app.modules.registry import get_module_registry
from app.services.module_workflow_service import build_workflow_state


def build_decision_log(
    workflow_state: dict[str, Any],
    stage_defs: list,
    project_context: dict[str, Any] | None = None,
    current_user_email: str | None = None,
    *,
    module_id: str = "",
    module_name: str = "",
    module_instance_id: str = "",
) -> dict[str, Any]:
    """Build current-state rows for a single module workflow."""
    decisions = build_current_state_rows(
        workflow_state=workflow_state,
        stage_defs=stage_defs,
        module_id=module_id,
        module_name=module_name,
        module_instance_id=module_instance_id,
    )
    return {
        "metadata": {
            "project_title": (project_context or {}).get("project_title", ""),
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "generated_by": current_user_email or "system",
            "total_decisions": len(decisions),
        },
        "decisions": decisions,
    }


def build_current_state_rows(
    *,
    workflow_state: dict[str, Any],
    stage_defs: list,
    module_id: str,
    module_name: str,
    module_instance_id: str,
) -> list[dict[str, Any]]:
    stage_id_to_title = {s.id: s.title for s in stage_defs}
    rows_with_sort_key: list[tuple[str, dict[str, Any]]] = []
    final_approval = workflow_state.get("final_approval") or {}
    attribution = _decision_log_attribution_for_module(module_id)

    for stage_id, stage_state in workflow_state.get("stages", {}).items():
        if stage_state.get("status") not in ("confirmed", "draft"):
            continue

        stage_title = stage_id_to_title.get(stage_id, stage_id)
        stage_data = stage_state.get("data") or {}
        confirmed_by = (
            stage_state.get("confirmed_by_email")
            or stage_state.get("confirmed_by")
            or "—"
        )
        confirmed_at = stage_state.get("confirmed_at") or ""
        status = stage_state.get("status") or "pending"

        for item in stage_data.get("items", []):
            item_rows = _build_item_rows(
                item=item,
                module_id=module_id,
                module_name=module_name,
                module_instance_id=module_instance_id,
                stage_id=stage_id,
                stage_title=stage_title,
                status=status,
                confirmed_by=confirmed_by,
                confirmed_at=confirmed_at,
                final_approval=final_approval,
                attribution=attribution,
            )
            rows_with_sort_key.extend((confirmed_at, row) for row in item_rows)

        for item_id, record in (stage_data.get("records") or {}).items():
            record_rows = _build_record_rows(
                record=record,
                item_id=item_id,
                module_id=module_id,
                module_name=module_name,
                module_instance_id=module_instance_id,
                stage_id=stage_id,
                stage_title=stage_title,
                status=status,
                confirmed_by=confirmed_by,
                confirmed_at=confirmed_at,
                final_approval=final_approval,
                attribution=attribution,
            )
            rows_with_sort_key.extend((confirmed_at, row) for row in record_rows)

        widget_data = stage_data.get("widget_data")
        if isinstance(widget_data, dict):
            widget_rows = _build_widget_rows(
                widget_data=widget_data,
                module_id=module_id,
                module_name=module_name,
                module_instance_id=module_instance_id,
                stage_id=stage_id,
                stage_title=stage_title,
                status=status,
                confirmed_by=confirmed_by,
                confirmed_at=confirmed_at,
                final_approval=final_approval,
                attribution=attribution,
            )
            rows_with_sort_key.extend((confirmed_at, row) for row in widget_rows)

    # Most recent first for UI and export chronology.
    rows_with_sort_key.sort(key=lambda pair: pair[0] or "", reverse=True)
    return [row for _, row in rows_with_sort_key]


def build_module_decision_history_report(
    *,
    workflow_state: dict[str, Any],
    stage_defs: list,
    module_id: str,
    module_name: str,
    module_instance_id: str,
) -> dict[str, Any]:
    """Build a module-scoped, value-level decision history report.

    History rows are value-centric entries (field/value/source/provenance)
    for the module workflow, including confirmation/final-approval metadata.
    """
    history_rows = build_current_state_rows(
        workflow_state=workflow_state,
        stage_defs=stage_defs,
        module_id=module_id,
        module_name=module_name,
        module_instance_id=module_instance_id,
    )
    return {
        "metadata": {
            "module_id": module_id,
            "module_name": module_name,
            "module_instance_id": module_instance_id,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "history_row_count": len(history_rows),
        },
        "history_rows": history_rows,
    }


async def build_initiative_decision_log(
    db: AsyncSession,
    *,
    initiative_id: UUID,
    module_instance_id: UUID | None = None,
    module_id: str | None = None,
) -> dict[str, Any]:
    initiative = await db.get(Initiative, initiative_id)
    stmt = select(ModuleInstance).where(
        ModuleInstance.initiative_id == initiative_id,
        ModuleInstance.archived.is_(False),
    )
    if module_instance_id is not None:
        stmt = stmt.where(ModuleInstance.id == module_instance_id)
    if module_id:
        stmt = stmt.where(ModuleInstance.module_id == module_id)
    stmt = stmt.order_by(ModuleInstance.updated_at.desc())
    instances = list((await db.execute(stmt)).scalars().all())

    registry = get_module_registry()
    current_rows: list[dict[str, Any]] = []
    instance_ids = [inst.id for inst in instances]

    for inst in instances:
        module = registry.get_module(inst.module_id)
        if module is None:
            continue
        state = await build_workflow_state(db, inst, module)
        current_rows.extend(
            build_current_state_rows(
                workflow_state=state,
                stage_defs=module.stage_defs,
                module_id=inst.module_id,
                module_name=module.definition.name,
                module_instance_id=str(inst.id),
            )
        )

    history_rows: list[dict[str, Any]] = []
    if instance_ids:
        event_stmt = select(DecisionEvent).where(
            DecisionEvent.module_instance_id.in_(instance_ids)
        ).order_by(DecisionEvent.created_at.desc(), DecisionEvent.sequence_number.desc())
        events = list((await db.execute(event_stmt)).scalars().all())
        history_rows = [_event_to_history_row(event) for event in events]

    return {
        "metadata": {
            "project_title": initiative.title if initiative else "",
            "initiative_id": str(initiative_id),
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "current_row_count": len(current_rows),
            "history_row_count": len(history_rows),
        },
        "current_rows": current_rows,
        "history_rows": history_rows,
    }


def build_decision_log_xlsx(report: dict[str, Any]) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    except ImportError as exc:
        raise RuntimeError("openpyxl is required for XLSX export") from exc

    wb = openpyxl.Workbook()
    current_sheet = wb.active
    current_sheet.title = "Current State"
    history_sheet = wb.create_sheet("History")

    header_fill = PatternFill("solid", fgColor="F3F4F6")
    border = Border(bottom=Side(style="thin", color="D1D5DB"))

    _write_sheet(
        current_sheet,
        report.get("current_rows", []),
        header_fill=header_fill,
        border=border,
        Font=Font,
        Alignment=Alignment,
    )
    _write_sheet(
        history_sheet,
        report.get("history_rows", []),
        header_fill=header_fill,
        border=border,
        Font=Font,
        Alignment=Alignment,
    )

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_module_decision_log_xlsx(report: dict[str, Any]) -> bytes:
    """Build a module-scoped decision history workbook (single History sheet)."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    except ImportError as exc:
        raise RuntimeError("openpyxl is required for XLSX export") from exc

    wb = openpyxl.Workbook()
    history_sheet = wb.active
    history_sheet.title = "History"

    header_fill = PatternFill("solid", fgColor="F3F4F6")
    border = Border(bottom=Side(style="thin", color="D1D5DB"))

    _write_sheet(
        history_sheet,
        report.get("history_rows", []),
        header_fill=header_fill,
        border=border,
        Font=Font,
        Alignment=Alignment,
    )

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_decision_log_docx(
    workflow_state: dict[str, Any],
    stage_defs: list,
    project_context: dict[str, Any] | None = None,
    current_user_email: str | None = None,
) -> bytes:
    """Build a module-scoped DOCX summary from current decision rows."""
    log = build_decision_log(workflow_state, stage_defs, project_context, current_user_email)
    return _render_docx(log)


def _build_item_rows(
    *,
    item: dict[str, Any],
    module_id: str,
    module_name: str,
    module_instance_id: str,
    stage_id: str,
    stage_title: str,
    status: str,
    confirmed_by: str,
    confirmed_at: str,
    final_approval: dict[str, Any],
    attribution: DecisionLogAttribution,
) -> list[dict[str, Any]]:
    content = item.get("content") or {}
    item_id = item.get("id") or ""
    item_label = _item_label(content)
    source_type, source_detail = _classify_item_source(item, attribution)
    current_value = _item_value(content)
    return [
        _base_row(
            module_id=module_id,
            module_name=module_name,
            module_instance_id=module_instance_id,
            stage_id=stage_id,
            stage_title=stage_title,
            entity_type="item",
            entity_id=item_id,
            item_label=item_label,
            field_label="Value",
            current_value=current_value,
            source_type=source_type,
            source_detail=source_detail,
            status=status,
            confirmed_by=confirmed_by,
            confirmed_at=confirmed_at,
            final_approval=final_approval,
        )
    ]


def _build_record_rows(
    *,
    record: dict[str, Any],
    item_id: str,
    module_id: str,
    module_name: str,
    module_instance_id: str,
    stage_id: str,
    stage_title: str,
    status: str,
    confirmed_by: str,
    confirmed_at: str,
    final_approval: dict[str, Any],
    attribution: DecisionLogAttribution,
) -> list[dict[str, Any]]:
    record_source_type, record_source_detail = _classify_record_source(record, attribution)
    rows: list[dict[str, Any]] = []
    for field_name, value in _flatten_scalars(record):
        rows.append(
            _base_row(
                module_id=module_id,
                module_name=module_name,
                module_instance_id=module_instance_id,
                stage_id=stage_id,
                stage_title=stage_title,
                entity_type="record",
                entity_id=item_id,
                item_label=f"Record {item_id[:8]}",
                field_label=field_name.replace(".", " / ").replace("_", " ").title(),
                current_value=value,
                source_type=record_source_type,
                source_detail=record_source_detail,
                status=status,
                confirmed_by=confirmed_by,
                confirmed_at=confirmed_at,
                final_approval=final_approval,
            )
        )
    return rows


def _build_widget_rows(
    *,
    widget_data: dict[str, Any],
    module_id: str,
    module_name: str,
    module_instance_id: str,
    stage_id: str,
    stage_title: str,
    status: str,
    confirmed_by: str,
    confirmed_at: str,
    final_approval: dict[str, Any],
    attribution: DecisionLogAttribution,
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    metrics = _computed_overview_metrics(module_id, widget_data)
    integration_citation = _computed_integration_citation(module_id, widget_data, attribution)
    for metric_label, metric_value in metrics:
        rows.append(
            _base_row(
                module_id=module_id,
                module_name=module_name,
                module_instance_id=module_instance_id,
                stage_id=stage_id,
                stage_title=stage_title,
                entity_type="computed_result",
                entity_id=metric_label.lower().replace(" ", "_"),
                item_label=metric_label,
                field_label="",
                current_value=metric_value,
                source_type="Computed Value",
                source_detail=integration_citation,
                status=status,
                confirmed_by=confirmed_by,
                confirmed_at=confirmed_at,
                final_approval=final_approval,
            )
        )
    return rows


def _computed_overview_metrics(module_id: str, widget_data: dict[str, Any]) -> list[tuple[str, Any]]:
    result = widget_data.get("result") if isinstance(widget_data.get("result"), dict) else {}
    if not result:
        # Safe fallback: only top-level scalar outputs, never nested input payloads.
        return [
            (key.replace("_", " ").title(), value)
            for key, value in widget_data.items()
            if key not in {"inputs", "sensitivity", "cash_flows", "er_schedule", "ac_monthly", "dc_monthly", "solrad_monthly", "poa_monthly"}
            and _has_scalar_value(value)
        ]

    if module_id == "lcoe_model":
        currency = result.get("currency") or "USD"
        raw_lcoe = result.get("lcoe")
        try:
            lcoe_value = float(raw_lcoe) if raw_lcoe is not None else None
        except (TypeError, ValueError):
            lcoe_value = None
        metrics = [
            ("LCOE", f"{currency} {lcoe_value:.4f} /kWh" if lcoe_value is not None else None),
            ("Discounted Costs (NPV)", _fmt_number_with_unit(result.get("npv_total_costs"), currency)),
            ("Discounted Energy (NPV)", _fmt_number_with_unit(result.get("npv_total_energy"), "kWh")),
            ("Total Production", _fmt_number_with_unit(result.get("lifetime_energy_kwh"), "kWh")),
            ("Assumptions Used", result.get("assumption_count")),
            ("Confidence", result.get("quality_label")),
        ]
        return [(label, value) for label, value in metrics if _has_scalar_value(value)]

    if module_id == "carbon_model":
        metrics = [
            ("Net ERs", _fmt_number_with_unit(result.get("net_er_tco2e"), "tCO2e/yr")),
            ("Baseline Emissions", _fmt_number_with_unit(result.get("baseline_emissions_tco2e"), "tCO2e/yr")),
            ("Project Emissions", _fmt_number_with_unit(result.get("project_emissions_tco2e"), "tCO2e/yr")),
            ("Leakage", _fmt_number_with_unit(result.get("leakage_tco2e"), "tCO2e/yr")),
            ("Crediting Period", _fmt_number_with_unit(result.get("period_years"), "years")),
            ("Assumptions Used", result.get("assumption_count")),
            ("Confidence", result.get("quality_label")),
        ]
        return [(label, value) for label, value in metrics if _has_scalar_value(value)]

    if module_id == "solar_estimate":
        metrics = [
            ("Annual AC Energy", _fmt_number_with_unit(result.get("ac_annual"), "kWh/yr")),
            ("Capacity Factor", _fmt_number_with_unit(result.get("capacity_factor"), "%")),
            ("Annual Solar Radiation", _fmt_number_with_unit(result.get("solrad_annual"), "kWh/m²/day")),
            ("Assumptions Used", widget_data.get("assumption_count")),
            ("Confidence", result.get("quality_label")),
        ]
        return [(label, value) for label, value in metrics if _has_scalar_value(value)]

    # Generic fallback for other computed widgets
    return [
        (key.replace("_", " ").title(), value)
        for key, value in result.items()
        if _has_scalar_value(value)
    ]


def _computed_integration_citation(
    module_id: str,
    widget_data: dict[str, Any],
    attribution: DecisionLogAttribution,
) -> str:
    citations: list[str] = []

    registry = get_module_registry()
    module = registry.get_module(module_id)
    manifest = module.manifest if module is not None else None
    adapter_bindings = manifest.adapter_bindings if manifest is not None else {}
    adapter_labels = attribution.adapter_labels
    widget_detail_labels = attribution.widget_detail_labels

    if attribution.include_adapter_bindings:
        for role, adapter_id in sorted(adapter_bindings.items()):
            citations.append(_adapter_citation(adapter_id, role, adapter_labels))

    result = widget_data.get("result") if isinstance(widget_data.get("result"), dict) else {}
    if attribution.include_provenance_sources:
        qualitative_refs = _source_references(_extract_source_list(widget_data) + _extract_source_list(result))
        if qualitative_refs:
            citations.append(qualitative_refs)

    if attribution.include_model_name:
        llm_model = _extract_model_name(widget_data) or _extract_model_name(result)
        if llm_model:
            citations.append(f"LLM model: {llm_model}")

    for key, label in sorted(widget_detail_labels.items()):
        value = widget_data.get(key)
        if _has_scalar_value(value):
            citations.append(f"{label}: {value}")

    unique: list[str] = []
    for citation in citations:
        if citation and citation not in unique:
            unique.append(citation)
    return "; ".join(unique[:4]) or "—"


def _adapter_citation(adapter_id: str, role: str, adapter_labels: dict[str, str] | None = None) -> str:
    label = (adapter_labels or {}).get(adapter_id)
    if label:
        return f"{label} ({role}: {adapter_id})"
    return f"Adapter ({role}): {adapter_id}"


def _decision_log_attribution_for_module(module_id: str) -> DecisionLogAttribution:
    registry = get_module_registry()
    module = registry.get_module(module_id)
    if module is None:
        return DecisionLogAttribution()
    return module.manifest.decision_log_attribution


def _classify_record_source(
    record: dict[str, Any],
    attribution: DecisionLogAttribution,
) -> tuple[str, str]:
    references = _source_references(_extract_source_list(record)) if attribution.include_provenance_sources else ""
    llm_model = _extract_model_name(record) if attribution.include_model_name else None
    if references:
        detail = references
    else:
        detail = "LLM enrichment or manual entry"
    if llm_model:
        detail = f"{detail}; model: {llm_model}"
    return ("LLM Enrichment", detail)


def _extract_source_list(value: Any) -> list[Any]:
    if not isinstance(value, dict):
        return []
    candidates: list[Any] = []
    for key in ("sources", "citations", "references"):
        source_list = value.get(key)
        if isinstance(source_list, list):
            candidates.extend(source_list)
    provenance = value.get("provenance")
    if isinstance(provenance, dict):
        nested_sources = provenance.get("sources")
        if isinstance(nested_sources, list):
            candidates.extend(nested_sources)
    return candidates


def _extract_model_name(value: Any) -> str | None:
    if not isinstance(value, dict):
        return None
    for key in ("llm_model", "model_name", "model"):
        model_name = value.get(key)
        if isinstance(model_name, str) and model_name.strip():
            return model_name.strip()
    provenance = value.get("provenance")
    if isinstance(provenance, dict):
        for key in ("llm_model", "model_name", "model"):
            model_name = provenance.get(key)
            if isinstance(model_name, str) and model_name.strip():
                return model_name.strip()
    metadata = value.get("metadata")
    if isinstance(metadata, dict):
        for key in ("llm_model", "model_name", "model"):
            model_name = metadata.get(key)
            if isinstance(model_name, str) and model_name.strip():
                return model_name.strip()
    return None


def _fmt_number_with_unit(value: Any, unit: str) -> str | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return f"{'1' if value else '0'} {unit}".strip()
    if isinstance(value, int):
        return f"{value:,} {unit}".strip()
    if isinstance(value, float):
        formatted = f"{value:,.2f}".rstrip("0").rstrip(".")
        return f"{formatted} {unit}".strip()
    return f"{value} {unit}".strip()


def _base_row(
    *,
    module_id: str,
    module_name: str,
    module_instance_id: str,
    stage_id: str,
    stage_title: str,
    entity_type: str,
    entity_id: str,
    item_label: str,
    field_label: str,
    current_value: Any,
    source_type: str,
    source_detail: str,
    status: str,
    confirmed_by: str,
    confirmed_at: str,
    final_approval: dict[str, Any],
) -> dict[str, Any]:
    return {
        "module": module_name or module_id,
        "module_id": module_id,
        "module_instance_id": module_instance_id,
        "stage": stage_title,
        "stage_id": stage_id,
        "entity_type": entity_type,
        "entity_id": entity_id,
        "item": item_label,
        "current_value": _stringify_value(current_value),
        "source_type": source_type,
        "source_detail": source_detail,
        "status": status,
        "confirmed_by": confirmed_by or "—",
        "confirmed_at": _format_ts(confirmed_at),
        "final_approved_by": (final_approval.get("approved_by_email") or final_approval.get("approved_by") or "—"),
        "final_approved_at": _format_ts(final_approval.get("approved_at") or ""),
    }


def _item_value(content: dict[str, Any]) -> Any:
    unit = content.get("unit")
    preferred_keys = ("value", "amount", "score", "answer", "selection", "selected")

    for key in preferred_keys:
        value = content.get(key)
        if _has_scalar_value(value):
            if unit and key in {"value", "amount", "score"}:
                return f"{value} {unit}"
            return value

    ignored_keys = {
        "status",
        "category",
        "variable",
        "field_name",
        "field_type",
        "options",
        "placeholder",
        "icon",
        "source",
        "unit",
    }
    for key, value in content.items():
        if key in ignored_keys:
            continue
        if not _has_scalar_value(value):
            continue
        return f"{value} {unit}" if unit and key == "value" else value

    return item_fallback_value(content)


def _has_scalar_value(value: Any) -> bool:
    if value is None or value == "":
        return False
    if isinstance(value, (dict, list, tuple, set)):
        return False
    return True


def item_fallback_value(content: dict[str, Any]) -> str:
    return str(
        content.get("name")
        or content.get("label")
        or content.get("title")
        or content.get("variable")
        or "—"
    )


def _flatten_scalars(value: Any, prefix: str = "") -> list[tuple[str, Any]]:
    rows: list[tuple[str, Any]] = []
    if isinstance(value, dict):
        for key, nested in value.items():
            next_prefix = f"{prefix}.{key}" if prefix else key
            rows.extend(_flatten_scalars(nested, next_prefix))
        return rows
    if isinstance(value, list):
        for idx, nested in enumerate(value):
            next_prefix = f"{prefix}[{idx}]"
            rows.extend(_flatten_scalars(nested, next_prefix))
        return rows
    if value is None or value == "":
        return rows
    return [(prefix or "value", value)]


def _item_label(content: dict[str, Any]) -> str:
    return str(
        content.get("name")
        or content.get("label")
        or content.get("variable")
        or content.get("title")
        or content.get("category")
        or "Item"
    )


def _classify_item_source(
    item: dict[str, Any],
    attribution: DecisionLogAttribution,
) -> tuple[str, str]:
    provenance = item.get("provenance") or {}
    derivation = str(provenance.get("derivation") or item.get("origin") or "").lower()
    sources = _extract_source_list(item) if attribution.include_provenance_sources else []
    references = _source_references(sources) if attribution.include_provenance_sources else ""
    llm_model = _extract_model_name(item) if attribution.include_model_name else None

    def _detail(base: str) -> str:
        if llm_model:
            if base == "—":
                return f"model: {llm_model}"
            return f"{base}; model: {llm_model}"
        return base

    if "user" in derivation or "provided" in derivation:
        return ("User Input", "—")
    if "template" in derivation:
        return ("Prior Module Output", _detail("—"))
    if any(_is_external_source(source) for source in sources):
        return ("External Research", _detail(references or "—"))
    if any(_is_internal_source(source) for source in sources):
        return ("Internal Materials", _detail(references or "—"))
    fallback = references or "Model-generated inference"
    return ("LLM Inference", _detail(fallback))


def _is_internal_source(source: dict[str, Any]) -> bool:
    source_type = str(source.get("source_type") or "").lower()
    return source_type in {"corpus", "evidence", "conversation"}


def _is_external_source(source: dict[str, Any]) -> bool:
    source_type = str(source.get("source_type") or "").lower()
    return source_type in {"web", "openalex", "url", "external"}


def _source_references(sources: list[Any]) -> str:
    references: list[str] = []
    for source in sources:
        if not isinstance(source, dict):
            continue
        title = source.get("source_title") or source.get("title")
        url = source.get("source_url") or source.get("url")
        source_type = str(source.get("source_type") or "").lower()
        if title and url:
            references.append(f"{title} ({url})")
        elif title:
            references.append(str(title))
        elif url:
            references.append(str(url))
        elif source_type:
            references.append(source_type)

    unique_references: list[str] = []
    for ref in references:
        if ref not in unique_references:
            unique_references.append(ref)
    return "; ".join(unique_references[:3])


def _event_to_history_row(event: DecisionEvent) -> dict[str, Any]:
    payload = event.payload_json or {}
    return {
        "module_id": event.module_id,
        "module_instance_id": str(event.module_instance_id),
        "stage_id": event.stage_id or "",
        "event": event.event_type.replace("_", " ").title(),
        "entity_type": event.entity_type.replace("_", " ").title(),
        "entity_id": event.entity_id or "—",
        "actor": event.actor_email or event.actor_user_id or "system",
        "occurred_at": _format_ts(event.created_at.isoformat()),
        "details": _summarize_payload(payload),
    }


def _summarize_payload(payload: dict[str, Any]) -> str:
    if not payload:
        return "—"
    keys = list(payload.keys())[:4]
    parts = [f"{key}: {_stringify_value(payload.get(key))}" for key in keys]
    return "; ".join(parts)


def _write_sheet(sheet, rows: list[dict[str, Any]], *, header_fill, border, Font, Alignment) -> None:
    if not rows:
        sheet["A1"] = "No rows available."
        return

    headers = list(rows[0].keys())
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(bold=True, size=10)
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal="left", vertical="center")

    for row in rows:
        sheet.append([row.get(header, "") for header in headers])

    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for idx, header in enumerate(headers, start=1):
        max_len = max(len(str(header)), *(len(str(r.get(header, ""))) for r in rows[:200]))
        sheet.column_dimensions[_column_letter(idx)].width = min(max(max_len + 2, 12), 36)


def _column_letter(idx: int) -> str:
    letters = ""
    while idx:
        idx, remainder = divmod(idx - 1, 26)
        letters = chr(65 + remainder) + letters
    return letters


def _format_ts(iso_str: str) -> str:
    if not iso_str:
        return "—"
    try:
        dt = datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
        return dt.strftime("%d %b %Y %H:%M UTC")
    except Exception:
        return iso_str


def _stringify_value(value: Any) -> str:
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if value is None:
        return "—"
    if isinstance(value, float):
        return f"{value:g}"
    return str(value)


def _render_docx(log: dict[str, Any]) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX export") from exc

    doc = Document()
    meta = log.get("metadata", {})
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(f"Decision Log — {meta.get('project_title', 'Project')}")
    title_run.bold = True
    title_run.font.size = Pt(16)

    meta_para = doc.add_paragraph(
        f"Generated: {_format_ts(meta.get('generated_at', ''))}  ·  By: {meta.get('generated_by', 'system')}  ·  {meta.get('total_decisions', 0)} rows"
    )
    meta_para.runs[0].font.size = Pt(9)
    meta_para.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()

    decisions = log.get("decisions", [])
    if not decisions:
        doc.add_paragraph("No decision rows found.")
    else:
        headers = ["Stage", "Item", "Field", "Value", "Source", "Confirmed By", "Confirmed At"]
        keys = ["stage", "item", "field", "current_value", "source_type", "confirmed_by", "confirmed_at"]
        widths = [1.2, 1.8, 1.5, 1.8, 1.4, 1.4, 1.3]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, (header, width) in enumerate(zip(headers, widths)):
            cell = table.rows[0].cells[idx]
            cell.width = Inches(width)
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            tc = cell._tc
            shd = tc.get_or_add_tcPr().get_or_add_shd()
            shd.set(qn("w:fill"), "005E72")
            shd.set(qn("w:color"), "auto")
            para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for decision in decisions:
            row_cells = table.add_row().cells
            for idx, (key, width) in enumerate(zip(keys, widths)):
                row_cells[idx].width = Inches(width)
                row_cells[idx].paragraphs[0].add_run(str(decision.get(key, ""))).font.size = Pt(8.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
