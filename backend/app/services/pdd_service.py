"""PDD (Project Design Document) authoring service.

Orchestrates the multi-step PDD creation flow:
scan -> outline -> section-by-section authoring -> consistency review -> assembly.
"""

from __future__ import annotations

import asyncio
import json
import logging
from datetime import datetime, timezone
from typing import Any
from uuid import UUID

from openai import AsyncOpenAI
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm.attributes import flag_modified

from app.config import get_settings
from app.models.initiative import Initiative
from app.models.pdd import PDDWorkspace
from app.models.evidence import EvidenceDoc
from app.models.project_material import ProjectMaterial
from app.services.rag import RAGService, RetrievedChunk

settings = get_settings()
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# LLM function-calling schemas
# ---------------------------------------------------------------------------

SCAN_PROJECT_SCHEMA = {
    "type": "function",
    "function": {
        "name": "scan_project",
        "description": "Analyse available project materials and return a structured scan.",
        "parameters": {
            "type": "object",
            "required": [
                "project_type",
                "project_type_label",
                "pdd_style",
                "sources_summary",
                "information_gaps",
            ],
            "properties": {
                "project_type": {
                    "type": "string",
                    "description": "Classified project type id (e.g. renewable_energy, clean_cooking, forestry, agriculture, waste_management, general).",
                },
                "project_type_label": {
                    "type": "string",
                    "description": "Human-readable label for the project type.",
                },
                "pdd_style": {
                    "type": "string",
                    "description": "Likely PDD structure or standard the project aligns to, if inferable (e.g. 'Gold Standard PDD', 'Verra VCS PDD', 'Custom'). Use 'Adaptive' when no single standard is clear.",
                },
                "sources_summary": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "filename": {"type": "string"},
                            "topics_covered": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                        },
                        "required": ["filename", "topics_covered"],
                    },
                    "description": "Summary of each source document and the topics it covers.",
                },
                "information_gaps": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Major information gaps visible from the available materials.",
                },
            },
        },
    },
}

GENERATE_OUTLINE_SCHEMA = {
    "type": "function",
    "function": {
        "name": "generate_outline",
        "description": "Generate a PDD outline adapted to the project type and available materials.",
        "parameters": {
            "type": "object",
            "required": ["sections"],
            "properties": {
                "sections": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "string"},
                            "title": {"type": "string"},
                            "description": {"type": "string"},
                            "key_topics": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                        },
                        "required": ["id", "title", "description", "key_topics"],
                    },
                },
            },
        },
    },
}

PREPARE_SECTION_SCHEMA = {
    "type": "function",
    "function": {
        "name": "prepare_section",
        "description": "Analyse evidence for a PDD section and identify gaps.",
        "parameters": {
            "type": "object",
            "required": ["evidence_notes", "missing_items", "follow_up_questions"],
            "properties": {
                "evidence_notes": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "citation_key": {"type": "integer"},
                            "note": {"type": "string"},
                        },
                        "required": ["citation_key", "note"],
                    },
                    "description": "Short notes explaining how each piece of evidence relates to this section.",
                },
                "missing_items": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Information that is needed for this section but not found in the evidence.",
                },
                "follow_up_questions": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "string"},
                            "question": {"type": "string"},
                            "why": {"type": "string"},
                        },
                        "required": ["id", "question", "why"],
                    },
                    "description": "Targeted questions to ask the user before drafting this section.",
                },
            },
        },
    },
}

DRAFT_SECTION_SCHEMA = {
    "type": "function",
    "function": {
        "name": "draft_section",
        "description": "Draft PDD section content with inline citations.",
        "parameters": {
            "type": "object",
            "required": ["content", "citations_used", "confidence", "unsupported_claims"],
            "properties": {
                "content": {
                    "type": "string",
                    "description": "Full section text with inline [N] citations referencing provided evidence.",
                },
                "citations_used": {
                    "type": "array",
                    "items": {"type": "integer"},
                    "description": "List of citation numbers actually referenced in the text.",
                },
                "confidence": {
                    "type": "string",
                    "enum": ["high", "medium", "low"],
                    "description": "Overall confidence in the drafted section based on evidence quality.",
                },
                "unsupported_claims": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Statements in the draft that are NOT backed by evidence (user should verify).",
                },
            },
        },
    },
}

CONSISTENCY_CHECK_SCHEMA = {
    "type": "function",
    "function": {
        "name": "consistency_check",
        "description": "Review all PDD sections for internal consistency.",
        "parameters": {
            "type": "object",
            "required": ["findings"],
            "properties": {
                "findings": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "string"},
                            "severity": {
                                "type": "string",
                                "enum": ["error", "warning", "info"],
                            },
                            "category": {
                                "type": "string",
                                "enum": [
                                    "terminology",
                                    "dates",
                                    "numbers",
                                    "geography",
                                    "implementation",
                                    "repeated_claims",
                                    "other",
                                ],
                            },
                            "description": {"type": "string"},
                            "affected_sections": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                            "suggestion": {"type": "string"},
                        },
                        "required": [
                            "id",
                            "severity",
                            "category",
                            "description",
                            "affected_sections",
                            "suggestion",
                        ],
                    },
                },
            },
        },
    },
}


# ---------------------------------------------------------------------------
# Service
# ---------------------------------------------------------------------------


class PDDService:
    """Orchestrates multi-step PDD authoring."""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.client = AsyncOpenAI(api_key=settings.openai_api_key)
        self.orchestration_model = settings.openai_orchestration_model
        self.generation_model = settings.openai_generation_model

    # -- helpers -------------------------------------------------------------

    async def _get_workspace(self, workspace_id: UUID) -> PDDWorkspace:
        ws = await self.db.get(PDDWorkspace, workspace_id)
        if ws is None:
            raise ValueError(f"PDDWorkspace {workspace_id} not found")
        return ws

    async def _get_workspace_by_initiative(self, initiative_id: UUID) -> PDDWorkspace | None:
        result = await self.db.execute(
            select(PDDWorkspace)
            .where(PDDWorkspace.initiative_id == initiative_id)
            .order_by(PDDWorkspace.created_at.desc())
            .limit(1)
        )
        return result.scalar_one_or_none()

    async def _get_initiative(self, initiative_id: UUID) -> Initiative:
        ini = await self.db.get(Initiative, initiative_id)
        if ini is None:
            raise ValueError(f"Initiative {initiative_id} not found")
        return ini

    async def _gather_material_texts(self, initiative_id: UUID) -> list[dict[str, str]]:
        """Return list of {filename, content_text} for all project materials + evidence."""
        materials: list[dict[str, str]] = []

        # Project materials with extracted text
        result = await self.db.execute(
            select(ProjectMaterial).where(ProjectMaterial.initiative_id == initiative_id)
        )
        for pm in result.scalars().all():
            if pm.content_text:
                materials.append({"filename": pm.filename, "content_text": pm.content_text[:6000]})

        # Evidence docs — get first chunk content as representative text
        result = await self.db.execute(
            select(EvidenceDoc).where(EvidenceDoc.initiative_id == initiative_id)
        )
        for ed in result.scalars().all():
            from app.models.evidence import EvidenceChunk
            chunk_result = await self.db.execute(
                select(EvidenceChunk)
                .where(EvidenceChunk.evidence_doc_id == ed.id)
                .order_by(EvidenceChunk.chunk_index)
                .limit(5)
            )
            chunks = chunk_result.scalars().all()
            if chunks:
                text = "\n".join(c.content for c in chunks)
                materials.append({"filename": ed.filename or "Evidence document", "content_text": text[:6000]})

        return materials

    def _touch(self, ws: PDDWorkspace) -> None:
        ws.updated_at = datetime.now(timezone.utc)
        flag_modified(ws, "updated_at")

    # -- 1. create workspace -------------------------------------------------

    async def create_workspace(
        self,
        initiative_id: UUID,
        session_id: UUID | None = None,
    ) -> PDDWorkspace:
        existing = await self._get_workspace_by_initiative(initiative_id)
        if existing is not None:
            return existing

        ws = PDDWorkspace(
            initiative_id=initiative_id,
            session_id=session_id,
            status="scan",
        )
        self.db.add(ws)
        await self.db.commit()
        await self.db.refresh(ws)
        return ws

    # -- 2. scan project -----------------------------------------------------

    async def scan_project(self, initiative_id: UUID) -> dict[str, Any]:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            ws = await self.create_workspace(initiative_id)

        initiative = await self._get_initiative(initiative_id)
        materials = await self._gather_material_texts(initiative_id)

        materials_block = ""
        for m in materials:
            materials_block += f"\n--- {m['filename']} ---\n{m['content_text']}\n"

        if not materials_block.strip():
            materials_block = "(No project materials uploaded yet.)"

        project_context = (
            f"Project title: {initiative.title or 'Untitled Project'}\n"
            f"Project description: {initiative.project_description or 'Not provided'}\n"
            f"Project type (auto-classified): {initiative.project_type or 'unknown'}\n"
            f"Geography: {initiative.geography or 'Not specified'}\n"
            f"Goal: {initiative.goal or 'Not specified'}\n"
        )

        response = await self.client.chat.completions.create(
            model=self.orchestration_model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a project design document specialist. Analyse the project context and "
                        "available materials and return a structured scan.\n\n"
                        "IMPORTANT: Base the project_type and project_type_label on the project title and "
                        "description provided — these are authoritative. Do NOT guess or contradict them. "
                        "If the title says 'Solar Farm', the project type is solar energy, not wind."
                    ),
                },
                {
                    "role": "user",
                    "content": f"{project_context}\n\nAvailable materials:\n{materials_block}",
                },
            ],
            tools=[SCAN_PROJECT_SCHEMA],
            tool_choice={"type": "function", "function": {"name": "scan_project"}},
            temperature=0.3,
        )

        tool_call = response.choices[0].message.tool_calls[0]
        scan_data = json.loads(tool_call.function.arguments)

        ws.project_scan = scan_data
        ws.status = "scan"
        self._touch(ws)
        flag_modified(ws, "project_scan")
        await self.db.commit()
        await self.db.refresh(ws)

        return scan_data

    # -- 3. generate outline -------------------------------------------------

    async def generate_outline(self, initiative_id: UUID) -> list[dict]:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            raise ValueError("No PDD workspace found. Run scan first.")

        initiative = await self._get_initiative(initiative_id)
        scan = ws.project_scan or {}

        response = await self.client.chat.completions.create(
            model=self.orchestration_model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a PDD outline specialist. Given the project scan results, propose a PDD outline "
                        "that is adapted to the project type and available materials.\n\n"
                        "Rules:\n"
                        "- Do NOT assume a single rigid PDD structure. Adapt based on the project type and scan.\n"
                        "- Include sections that are standard for this type of project.\n"
                        "- Each section needs a clear id, title, description, and 3-5 key topics.\n"
                        "- Typical PDD sections include: project overview, stakeholder analysis, "
                        "baseline scenario, project scenario, monitoring plan, risk assessment, etc.\n"
                        "- But adapt to the context — a forestry project differs from a clean cooking project.\n"
                        "- Aim for 8-15 sections."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Project title: {initiative.title or 'Untitled Project'}\n"
                        f"Project description: {initiative.project_description or 'No description'}\n"
                        f"Type: {scan.get('project_type_label', 'Unknown')}\n"
                        f"PDD style: {scan.get('pdd_style', 'Adaptive')}\n"
                        f"Available sources: {json.dumps(scan.get('sources_summary', []))}\n"
                        f"Known gaps: {json.dumps(scan.get('information_gaps', []))}\n"
                    ),
                },
            ],
            tools=[GENERATE_OUTLINE_SCHEMA],
            tool_choice={"type": "function", "function": {"name": "generate_outline"}},
            temperature=0.5,
        )

        tool_call = response.choices[0].message.tool_calls[0]
        outline_data = json.loads(tool_call.function.arguments)
        sections = outline_data.get("sections", [])

        ws.outline = sections
        ws.status = "outline"
        self._touch(ws)
        flag_modified(ws, "outline")
        await self.db.commit()
        await self.db.refresh(ws)

        return sections

    # -- 4. update outline (user edits) --------------------------------------

    async def update_outline(self, initiative_id: UUID, sections: list[dict]) -> list[dict]:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            raise ValueError("No PDD workspace found.")

        ws.outline = sections
        self._touch(ws)
        flag_modified(ws, "outline")
        await self.db.commit()
        await self.db.refresh(ws)
        return sections

    # -- 5. confirm outline --------------------------------------------------

    async def confirm_outline(self, initiative_id: UUID) -> dict:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            raise ValueError("No PDD workspace found.")

        outline = ws.outline or []
        if not outline:
            raise ValueError("Outline is empty.")

        # Initialise section state
        sections_state: dict[str, Any] = {}
        for section in outline:
            sid = section["id"]
            sections_state[sid] = {
                "status": "pending",
                "evidence": [],
                "missing_items": [],
                "questions": [],
                "draft": None,
                "citations": [],
                "confidence": None,
                "unsupported_claims": [],
                "user_answers": {},
            }

        first_section_id = outline[0]["id"] if outline else None

        ws.sections = sections_state
        ws.active_section_id = first_section_id
        ws.status = "authoring"
        self._touch(ws)
        flag_modified(ws, "sections")
        await self.db.commit()
        await self.db.refresh(ws)

        return {"status": "authoring", "active_section_id": first_section_id, "total_sections": len(outline)}

    # -- 6. prepare section (evidence + gaps) --------------------------------

    async def prepare_section(self, initiative_id: UUID, section_id: str) -> dict:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            raise ValueError("No PDD workspace found.")

        outline = ws.outline or []
        section_meta = next((s for s in outline if s["id"] == section_id), None)
        if section_meta is None:
            raise ValueError(f"Section '{section_id}' not in outline.")

        # RAG retrieval
        rag = RAGService(self.db)
        query = f"{section_meta['title']}: {section_meta['description']}. Key topics: {', '.join(section_meta.get('key_topics', []))}"
        chunks = await rag.retrieve(
            query=query,
            initiative_id=initiative_id,
            sources=["evidence", "corpus"],
            evidence_top_k=5,
            corpus_top_k=3,
        )

        # Build citation map
        citation_map: dict[str, int] = {}
        evidence_context = ""
        for i, chunk in enumerate(chunks):
            cnum = i + 1
            citation_map[str(chunk.chunk_id)] = cnum
            src_label = "[EVIDENCE]" if chunk.source_type == "evidence" else "[CORPUS]"
            evidence_context += f"\n[{cnum}] {src_label} {chunk.source_title}:\n{chunk.content}\n"

        if not evidence_context.strip():
            evidence_context = "(No relevant evidence found.)"

        # LLM analysis
        response = await self.client.chat.completions.create(
            model=self.orchestration_model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are helping author a Project Design Document. "
                        "Analyse the evidence provided for this section and identify:\n"
                        "1. How each piece of evidence relates to the section (evidence_notes)\n"
                        "2. What information is missing (missing_items)\n"
                        "3. Targeted follow-up questions to ask the user (follow_up_questions)\n\n"
                        "Be specific and practical. Only flag genuinely missing information."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Section: {section_meta['title']}\n"
                        f"Description: {section_meta['description']}\n"
                        f"Key topics: {', '.join(section_meta.get('key_topics', []))}\n\n"
                        f"Available evidence:\n{evidence_context}"
                    ),
                },
            ],
            tools=[PREPARE_SECTION_SCHEMA],
            tool_choice={"type": "function", "function": {"name": "prepare_section"}},
            temperature=0.3,
        )

        tool_call = response.choices[0].message.tool_calls[0]
        prep_data = json.loads(tool_call.function.arguments)

        # Build serialisable evidence list
        evidence_list = [
            {
                "citation_key": citation_map[str(c.chunk_id)],
                "chunk_id": str(c.chunk_id),
                "source_type": c.source_type,
                "source_title": c.source_title,
                "excerpt": c.content[:500],
                "similarity": round(c.similarity, 3),
            }
            for c in chunks
        ]

        # Persist on workspace
        sections = dict(ws.sections or {})
        sec = sections.get(section_id, {})
        sec["evidence"] = evidence_list
        sec["missing_items"] = prep_data.get("missing_items", [])
        sec["questions"] = prep_data.get("follow_up_questions", [])
        sec["evidence_notes"] = prep_data.get("evidence_notes", [])
        sec["status"] = "prepared"
        sections[section_id] = sec

        ws.sections = sections
        ws.active_section_id = section_id
        self._touch(ws)
        flag_modified(ws, "sections")
        await self.db.commit()
        await self.db.refresh(ws)

        return {
            "section_id": section_id,
            "evidence": evidence_list,
            "evidence_notes": prep_data.get("evidence_notes", []),
            "missing_items": prep_data.get("missing_items", []),
            "questions": prep_data.get("follow_up_questions", []),
        }

    # -- 7. draft section ----------------------------------------------------

    async def draft_section(
        self,
        initiative_id: UUID,
        section_id: str,
        user_answers: dict[str, str] | None = None,
        general_guidance: bool = False,
    ) -> dict:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            raise ValueError("No PDD workspace found.")

        outline = ws.outline or []
        section_meta = next((s for s in outline if s["id"] == section_id), None)
        if section_meta is None:
            raise ValueError(f"Section '{section_id}' not in outline.")

        sections = dict(ws.sections or {})
        sec = sections.get(section_id, {})

        # Build evidence context from stored evidence
        evidence_list = sec.get("evidence", []) if not general_guidance else []
        evidence_context = ""
        for ev in evidence_list:
            evidence_context += f"\n[{ev['citation_key']}] [{ev['source_type'].upper()}] {ev['source_title']}:\n{ev['excerpt']}\n"

        if not evidence_context.strip():
            evidence_context = (
                "(No project-specific evidence available. Draft using standard "
                "PDD content for this section type, flagging all claims as requiring verification.)"
                if general_guidance
                else "(No evidence available.)"
            )

        # Build user answers context
        answers_context = ""
        if user_answers:
            sec["user_answers"] = user_answers
            for qid, answer in user_answers.items():
                answers_context += f"- {qid}: {answer}\n"

        valid_citations = [ev["citation_key"] for ev in evidence_list]
        citation_list = ", ".join(f"[{n}]" for n in sorted(valid_citations)) if valid_citations else "none"

        system_content = (
            "You are drafting a section of a Project Design Document. "
            "Write clear, professional prose.\n\n"
        )
        if general_guidance:
            system_content += (
                "No project documents are available. Draft this section using general PDD best practices "
                "and standard content for this section type. Every factual claim MUST be listed in "
                "`unsupported_claims` since there is no evidence to back it up. Use placeholder text "
                "like '[TO BE CONFIRMED]' for project-specific values (names, numbers, locations).\n\n"
                "CITATION RULES:\n- Do NOT use any citations — there is no evidence.\n"
                "- Set confidence to 'low'."
            )
        else:
            system_content += (
                "Ground your writing in the provided evidence.\n\n"
                "CITATION RULES:\n"
                f"- You may ONLY use these citation numbers: {citation_list}\n"
                "- Use [N] inline when a claim is backed by evidence.\n"
                "- Do NOT invent citations.\n"
                "- If making a claim without evidence, do NOT add a citation.\n\n"
                "Mark any statements that are assumptions or require verification."
            )

        response = await self.client.chat.completions.create(
            model=self.generation_model,
            messages=[
                {"role": "system", "content": system_content},
                {
                    "role": "user",
                    "content": (
                        f"Section: {section_meta['title']}\n"
                        f"Description: {section_meta.get('description', '')}\n"
                        f"Key topics: {', '.join(section_meta.get('key_topics', []))}\n\n"
                        f"Evidence:\n{evidence_context}\n"
                        + (f"User-provided answers:\n{answers_context}\n" if answers_context else "")
                        + f"Missing items flagged: {json.dumps(sec.get('missing_items', []))}\n\n"
                        "Draft this section now."
                    ),
                },
            ],
            tools=[DRAFT_SECTION_SCHEMA],
            tool_choice={"type": "function", "function": {"name": "draft_section"}},
            temperature=0.6,
        )

        tool_call = response.choices[0].message.tool_calls[0]
        draft_data = json.loads(tool_call.function.arguments)

        sec["draft"] = draft_data.get("content", "")
        sec["citations"] = draft_data.get("citations_used", [])
        sec["confidence"] = draft_data.get("confidence", "medium")
        sec["unsupported_claims"] = draft_data.get("unsupported_claims", [])
        sec["status"] = "drafted"
        sections[section_id] = sec

        ws.sections = sections
        self._touch(ws)
        flag_modified(ws, "sections")
        await self.db.commit()
        await self.db.refresh(ws)

        return {
            "section_id": section_id,
            "content": draft_data.get("content", ""),
            "citations_used": draft_data.get("citations_used", []),
            "confidence": draft_data.get("confidence", "medium"),
            "unsupported_claims": draft_data.get("unsupported_claims", []),
        }

    # -- 8. update section (user edits) --------------------------------------

    async def update_section(self, initiative_id: UUID, section_id: str, content: str) -> None:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            raise ValueError("No PDD workspace found.")

        sections = dict(ws.sections or {})
        sec = sections.get(section_id)
        if sec is None:
            raise ValueError(f"Section '{section_id}' not found.")

        sec["draft"] = content
        sections[section_id] = sec
        ws.sections = sections
        self._touch(ws)
        flag_modified(ws, "sections")
        await self.db.commit()

    # -- 9. confirm section --------------------------------------------------

    async def confirm_section(self, initiative_id: UUID, section_id: str) -> dict:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            raise ValueError("No PDD workspace found.")

        sections = dict(ws.sections or {})
        sec = sections.get(section_id)
        if sec is None:
            raise ValueError(f"Section '{section_id}' not found.")

        sec["status"] = "confirmed"
        sections[section_id] = sec

        # Advance to next unconfirmed section
        outline = ws.outline or []
        next_section_id = None
        for s in outline:
            sid = s["id"]
            if sections.get(sid, {}).get("status") != "confirmed":
                next_section_id = sid
                break

        all_confirmed = next_section_id is None

        ws.sections = sections
        ws.active_section_id = next_section_id
        if all_confirmed:
            ws.status = "review"
        self._touch(ws)
        flag_modified(ws, "sections")
        await self.db.commit()
        await self.db.refresh(ws)

        return {
            "section_id": section_id,
            "next_section_id": next_section_id,
            "all_confirmed": all_confirmed,
        }

    # -- 10. consistency check -----------------------------------------------

    async def run_consistency_check(self, initiative_id: UUID) -> list[dict]:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            raise ValueError("No PDD workspace found.")

        outline = ws.outline or []
        sections = ws.sections or {}

        # Build full document text for review
        doc_text = ""
        for s in outline:
            sid = s["id"]
            sec = sections.get(sid, {})
            draft = sec.get("draft", "")
            if draft:
                doc_text += f"\n\n## {s['title']}\n{draft}"

        if not doc_text.strip():
            return []

        response = await self.client.chat.completions.create(
            model=self.orchestration_model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are reviewing a Project Design Document for internal consistency. "
                        "Check for contradictions or inconsistencies across sections in:\n"
                        "- Terminology (same concept named differently)\n"
                        "- Dates and timelines\n"
                        "- Numbers and quantities\n"
                        "- Geography and location references\n"
                        "- Implementation details\n"
                        "- Repeated or contradictory claims\n\n"
                        "Only flag genuine issues. Do not flag stylistic preferences."
                    ),
                },
                {"role": "user", "content": f"Full PDD draft:\n{doc_text}"},
            ],
            tools=[CONSISTENCY_CHECK_SCHEMA],
            tool_choice={"type": "function", "function": {"name": "consistency_check"}},
            temperature=0.2,
        )

        tool_call = response.choices[0].message.tool_calls[0]
        check_data = json.loads(tool_call.function.arguments)
        findings = check_data.get("findings", [])

        ws.consistency_findings = findings
        ws.status = "review"
        self._touch(ws)
        flag_modified(ws, "consistency_findings")
        await self.db.commit()
        await self.db.refresh(ws)

        return findings

    # -- 11. assemble document -----------------------------------------------

    async def assemble_document(self, initiative_id: UUID) -> dict:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            raise ValueError("No PDD workspace found.")

        initiative = await self._get_initiative(initiative_id)
        outline = ws.outline or []
        sections = ws.sections or {}

        # Build assembled sections
        assembled_sections = []
        all_citations: dict[int, dict] = {}
        unresolved_gaps: list[str] = []

        for s in outline:
            sid = s["id"]
            sec = sections.get(sid, {})
            draft = sec.get("draft", "")
            confidence = sec.get("confidence", "medium")
            unsupported = sec.get("unsupported_claims", [])
            missing = sec.get("missing_items", [])

            assembled_sections.append({
                "id": sid,
                "title": s["title"],
                "content": draft,
                "confidence": confidence,
                "unsupported_claims": unsupported,
            })

            # Collect citations
            for ev in sec.get("evidence", []):
                ckey = ev.get("citation_key")
                if ckey and ckey not in all_citations:
                    all_citations[ckey] = {
                        "number": ckey,
                        "source_type": ev.get("source_type", "evidence"),
                        "source_title": ev.get("source_title", "Unknown"),
                        "excerpt": ev.get("excerpt", ""),
                    }

            # Collect unresolved gaps
            for item in missing:
                if isinstance(item, str):
                    unresolved_gaps.append(f"[{s['title']}] {item}")

            for claim in unsupported:
                if isinstance(claim, str):
                    unresolved_gaps.append(f"[{s['title']}] Unsupported: {claim}")

        assembled = {
            "title": f"Project Design Document: {initiative.title or 'Untitled Project'}",
            "project_type": (ws.project_scan or {}).get("project_type_label", ""),
            "sections": assembled_sections,
            "citations": sorted(all_citations.values(), key=lambda c: c["number"]),
            "unresolved_gaps": unresolved_gaps,
            "section_count": len(assembled_sections),
            "citation_count": len(all_citations),
        }

        ws.assembled_document = assembled
        ws.missing_items_global = unresolved_gaps
        ws.status = "assembled"
        self._touch(ws)
        flag_modified(ws, "assembled_document")
        flag_modified(ws, "missing_items_global")
        await self.db.commit()
        await self.db.refresh(ws)

        return assembled

    # -- 12. export ----------------------------------------------------------

    async def export_docx(self, initiative_id: UUID) -> bytes:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None or ws.assembled_document is None:
            raise ValueError("No assembled PDD to export.")

        from docx import Document

        doc = Document()
        assembled = ws.assembled_document

        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        title_para = doc.add_heading(assembled.get("title", "Project Design Document"), level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        for section in assembled.get("sections", []):
            doc.add_heading(section["title"], level=1)
            content = section.get("content", "")
            for paragraph_text in content.split("\n\n"):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

        # Citations appendix
        citations = assembled.get("citations", [])
        if citations:
            doc.add_page_break()
            doc.add_heading("References", level=1)
            for cit in citations:
                doc.add_paragraph(
                    f"[{cit['number']}] {cit['source_title']} ({cit['source_type']}): {cit.get('excerpt', '')[:200]}...",
                    style="List Number",
                )

        # Unresolved gaps
        gaps = assembled.get("unresolved_gaps", [])
        if gaps:
            doc.add_heading("Unresolved Items", level=1)
            for gap in gaps:
                doc.add_paragraph(gap, style="List Bullet")

        import io

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # -- workspace state accessor -------------------------------------------

    async def get_workspace(self, initiative_id: UUID) -> dict | None:
        ws = await self._get_workspace_by_initiative(initiative_id)
        if ws is None:
            return None
        return {
            "id": str(ws.id),
            "initiative_id": str(ws.initiative_id) if ws.initiative_id else None,
            "status": ws.status,
            "project_scan": ws.project_scan,
            "outline": ws.outline,
            "sections": ws.sections,
            "active_section_id": ws.active_section_id,
            "consistency_findings": ws.consistency_findings,
            "assembled_document": ws.assembled_document,
            "missing_items_global": ws.missing_items_global,
            "created_at": ws.created_at.isoformat() if ws.created_at else None,
            "updated_at": ws.updated_at.isoformat() if ws.updated_at else None,
        }
