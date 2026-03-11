"""
Gold Standard Template Service

Handles fetching official GS templates from the Gold Standard website,
parsing DOCX structure to extract editable fields, caching results,
and managing template versions with an approval gate.
"""

import asyncio
import hashlib
import io
import logging
import re
import time
from collections import defaultdict
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
from typing import Optional
from uuid import UUID

import httpx
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.models.gs_template import GSTemplateVersion

settings = get_settings()
logger = logging.getLogger(__name__)

TEMPLATE_TYPE_COVER_LETTER = "cover_letter"
TEMPLATE_TYPE_PRELIMINARY_REVIEW = "preliminary_review"

TEMPLATE_URLS = {
    TEMPLATE_TYPE_COVER_LETTER: settings.gs_cover_letter_template_url,
    TEMPLATE_TYPE_PRELIMINARY_REVIEW: getattr(
        settings, "gs_preliminary_review_template_url", ""
    ) or None,
}


@dataclass
class FieldDef:
    field_id: str
    label: str
    field_type: str  # text, date, signature, multiline, choice
    section: str
    required: bool
    docx_location: dict  # {type: "paragraph"|"table_cell"|"content_control", index: ..., ...}
    placeholder_text: str = ""
    help_text: str = ""


@dataclass
class TemplateUpdateStatus:
    up_to_date: bool
    latest_approved_version_id: Optional[str] = None
    latest_approved_label: Optional[str] = None
    draft_available: bool = False
    draft_version_id: Optional[str] = None


# ---- In-memory TTL cache for update checks ----

_update_check_cache: dict[str, tuple[float, TemplateUpdateStatus]] = {}


class GSTemplateService:
    def __init__(self, db: AsyncSession):
        self.db = db

    # ------------------------------------------------------------------
    # Fetching
    # ------------------------------------------------------------------

    async def fetch_latest_template(self, template_type: str) -> GSTemplateVersion:
        """Download the DOCX from the GS website, compare hash, create draft if new."""
        url = TEMPLATE_URLS.get(template_type)
        if not url:
            raise ValueError(f"No known URL for template type '{template_type}'")

        docx_bytes = await self._download_template(url)
        file_hash = hashlib.sha256(docx_bytes).hexdigest()

        existing = await self._get_version_by_hash(template_type, file_hash)
        if existing:
            logger.info("Template %s hash %s already cached (id=%s)", template_type, file_hash[:12], existing.id)
            return existing

        logger.info("New template detected for %s (hash=%s), creating draft", template_type, file_hash[:12])
        field_schema = self._parse_template_fields(docx_bytes)
        html_preview = self._generate_html_preview(docx_bytes, field_schema)

        version_count = await self._count_versions(template_type)
        version_label = f"v{version_count + 1}"

        version = GSTemplateVersion(
            template_type=template_type,
            version_label=version_label,
            source_url=url,
            file_hash=file_hash,
            file_bytes=docx_bytes,
            html_preview=html_preview,
            field_schema=[asdict(f) for f in field_schema],
            status="draft",
        )
        self.db.add(version)
        await self.db.flush()
        await self.db.commit()
        await self.db.refresh(version)

        _update_check_cache.pop(template_type, None)
        return version

    async def check_for_updates(self, template_type: str) -> TemplateUpdateStatus:
        """Check whether the remote template has changed, with TTL caching.

        The remote network check is performed only when the TTL has expired AND
        we already have a version cached. If the network call fails or times out
        we return the last known status rather than blocking the caller.
        """
        cache_entry = _update_check_cache.get(template_type)
        ttl_seconds = settings.gs_template_cache_ttl_hours * 3600
        if cache_entry and (time.time() - cache_entry[0]) < ttl_seconds:
            return cache_entry[1]

        approved = await self._get_latest_by_status(template_type, "approved")
        draft = await self._get_latest_by_status(template_type, "draft")

        # Only do a live check if we already have something cached; never block cold start
        remote_changed = False
        if approved or draft:
            url = TEMPLATE_URLS.get(template_type)
            if url:
                try:
                    docx_bytes = await asyncio.wait_for(
                        self._download_template(url), timeout=15.0
                    )
                    remote_hash = hashlib.sha256(docx_bytes).hexdigest()
                    latest_hash = (approved.file_hash if approved else
                                   draft.file_hash if draft else None)
                    remote_changed = remote_hash != latest_hash
                    if remote_changed:
                        existing = await self._get_version_by_hash(template_type, remote_hash)
                        if not existing:
                            await self.fetch_latest_template(template_type)
                            draft = await self._get_latest_by_status(template_type, "draft")
                except (asyncio.TimeoutError, Exception):
                    logger.warning("Could not reach GS website to check for updates", exc_info=True)

        status = TemplateUpdateStatus(
            up_to_date=not remote_changed,
            latest_approved_version_id=str(approved.id) if approved else None,
            latest_approved_label=approved.version_label if approved else None,
            draft_available=draft is not None,
            draft_version_id=str(draft.id) if draft else None,
        )
        _update_check_cache[template_type] = (time.time(), status)
        return status

    async def get_active_template(self, template_type: str) -> Optional[GSTemplateVersion]:
        """Return the latest approved version, or the latest draft for bootstrap."""
        approved = await self._get_latest_by_status(template_type, "approved")
        if approved:
            return approved
        return await self._get_latest_by_status(template_type, "draft")

    async def get_or_fetch_active_template(self, template_type: str) -> GSTemplateVersion:
        """Return active template.

        Priority:
        1. Any approved version in the DB.
        2. Any draft version in the DB.
        3. Attempt a background-safe fetch from the GS website (short timeout).
        4. If the network call fails, create a synthetic template from fallback fields
           so the user can immediately start filling fields without being blocked.
        """
        active = await self.get_active_template(template_type)
        if active:
            return active

        # Nothing cached — try to fetch (skip if no URL), or create synthetic fallback
        url = TEMPLATE_URLS.get(template_type)
        if url:
            try:
                return await self.fetch_latest_template(template_type)
            except Exception as e:
                logger.warning(
                    "Could not fetch GS template from web (%s); creating synthetic fallback: %s",
                    template_type, e,
                )
        return await self._create_synthetic_template(template_type)

    async def approve_template(self, version_id: UUID, approved_by: str) -> GSTemplateVersion:
        """Approve a draft template and deprecate the previous approved version."""
        result = await self.db.execute(
            select(GSTemplateVersion).where(GSTemplateVersion.id == version_id)
        )
        version = result.scalar_one_or_none()
        if not version:
            raise ValueError(f"Template version {version_id} not found")
        if version.status != "draft":
            raise ValueError(f"Only draft templates can be approved (current: {version.status})")

        old_approved = await self._get_latest_by_status(version.template_type, "approved")
        if old_approved:
            old_approved.status = "deprecated"

        version.status = "approved"
        version.approved_by = approved_by
        version.approved_at = datetime.now(timezone.utc)
        await self.db.commit()
        await self.db.refresh(version)

        _update_check_cache.pop(version.template_type, None)
        return version

    async def get_version(self, version_id: UUID) -> Optional[GSTemplateVersion]:
        result = await self.db.execute(
            select(GSTemplateVersion).where(GSTemplateVersion.id == version_id)
        )
        return result.scalar_one_or_none()

    # ------------------------------------------------------------------
    # DOCX parsing
    # ------------------------------------------------------------------

    def _parse_template_fields(self, docx_bytes: bytes) -> list[FieldDef]:
        """Walk the DOCX body children in document order so section headings
        correctly apply to the tables that follow them."""
        doc = DocxDocument(io.BytesIO(docx_bytes))
        fields: list[FieldDef] = []
        seen_ids: set[str] = set()

        current_section = "General"
        para_index = 0
        table_index = 0

        for child in doc.element.body:
            tag = child.tag.split('}')[-1]

            if tag == 'p':
                if para_index < len(doc.paragraphs):
                    paragraph = doc.paragraphs[para_index]
                    text = paragraph.text.strip()
                    style_name = (paragraph.style.name
                                  if paragraph.style and paragraph.style.name else "")

                    if self._is_section_marker(text, style_name):
                        candidate = text.rstrip(':').strip()
                        if candidate and len(candidate) < 80:
                            current_section = (candidate.title()
                                               if candidate == candidate.upper()
                                               else candidate)

                    extracted = self._extract_fields_from_text(
                        text, current_section,
                        {"type": "paragraph", "index": para_index},
                    )
                    for f in extracted:
                        if f.field_id not in seen_ids:
                            fields.append(f)
                            seen_ids.add(f.field_id)
                para_index += 1

            elif tag == 'tbl':
                if table_index < len(doc.tables):
                    table = doc.tables[table_index]
                    table_fields = self._extract_fields_from_table(
                        table, table_index, current_section, seen_ids,
                    )
                    for f in table_fields:
                        fields.append(f)
                        seen_ids.add(f.field_id)
                table_index += 1

        sdt_fields = self._extract_content_controls(doc, current_section)
        for f in sdt_fields:
            if f.field_id not in seen_ids:
                fields.append(f)
                seen_ids.add(f.field_id)

        if not fields:
            fields = self._get_fallback_cover_letter_fields()

        return fields

    # ------------------------------------------------------------------

    @staticmethod
    def _is_section_marker(text: str, style_name: str) -> bool:
        """Heuristic to detect section headers."""
        if not text or len(text) < 3:
            return False
        if style_name.startswith("Heading"):
            level = style_name.replace("Heading", "").strip()
            if level.isdigit() and int(level) >= 6:
                return False
            return True
        if len(text) > 100 or text.startswith('['):
            return False
        if all(c in '_.… \u2026' for c in text):
            return False
        stripped = text.rstrip(':').strip()
        if stripped == stripped.upper() and len(stripped) > 3:
            return True
        if (len(text) < 60 and text[0].isupper()
                and not any(text.startswith(s)
                            for s in ('I ', 'The ', 'By ', 'Any ', 'A ', 'An '))):
            if len(text.split()) >= 2:
                return True
        return False

    def _extract_fields_from_text(self, text: str, section: str, location: dict) -> list[FieldDef]:
        """Find placeholder patterns in paragraph text."""
        fields: list[FieldDef] = []
        patterns = [
            (r'\[([^\]]+)\]', 'bracket'),       # [Enter project title]
            (r'\{\{([^}]+)\}\}', 'mustache'),    # {{project_title}}
            (r'Click here to enter (.+)', 'click_here'),
        ]
        for pattern, _style in patterns:
            for match in re.finditer(pattern, text):
                placeholder = match.group(1).strip()
                field_id = self._to_field_id(placeholder)
                label = self._clean_label(placeholder)
                field_type = self._infer_field_type(placeholder)
                fields.append(FieldDef(
                    field_id=field_id,
                    label=label,
                    field_type=field_type,
                    section=section,
                    required=True,
                    docx_location=location,
                    placeholder_text=match.group(0),
                ))
        return fields

    def _extract_fields_from_table(
        self, table, table_index: int, section: str, global_seen_ids: set[str],
    ) -> list[FieldDef]:
        """Extract bracket-placeholder fields from ALL cells in a table,
        handling merged cells and "LABEL: [placeholder]" patterns."""
        fields: list[FieldDef] = []
        local_seen: set[str] = set()

        for r_idx, row in enumerate(table.rows):
            processed_texts: set[str] = set()

            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()

                if cell_text in processed_texts:
                    continue
                processed_texts.add(cell_text)

                if len(cell_text) > 200 and '[' not in cell_text:
                    continue

                for match in re.finditer(r'\[([^\]]+)\]', cell_text):
                    placeholder = match.group(1).strip()
                    placeholder_full = match.group(0)

                    before = cell_text[:match.start()].strip()
                    before = re.sub(r'[:\.\s]+$', '', before).strip()
                    before = re.sub(r'\[.*?\]', '', before).strip()
                    before = re.sub(r'[:\.\s]+$', '', before).strip()

                    if before:
                        label_raw = before
                    elif c_idx > 0:
                        first_cell = row.cells[0].text.strip()
                        candidate = first_cell.split('[')[0].strip().rstrip(':').strip()
                        candidate_id = self._to_field_id(candidate) if candidate else ""
                        all_seen = global_seen_ids | local_seen
                        if candidate and candidate_id not in all_seen:
                            label_raw = candidate
                        else:
                            label_raw = placeholder
                    else:
                        label_raw = placeholder

                    if not label_raw:
                        label_raw = placeholder

                    label = self._clean_label(label_raw)
                    base_id = self._to_field_id(label_raw)
                    field_id = base_id
                    all_seen = global_seen_ids | local_seen

                    if field_id in all_seen:
                        section_prefix = self._to_field_id(section)
                        field_id = f"{section_prefix}_{base_id}"
                        counter = 2
                        while field_id in all_seen:
                            field_id = f"{section_prefix}_{base_id}_{counter}"
                            counter += 1

                    local_seen.add(field_id)
                    fields.append(FieldDef(
                        field_id=field_id,
                        label=label,
                        field_type=self._infer_field_type(label_raw),
                        section=section,
                        required=True,
                        docx_location={
                            "type": "table_cell",
                            "table": table_index,
                            "row": r_idx,
                            "col": c_idx,
                        },
                        placeholder_text=placeholder_full,
                    ))

        return fields

    def _extract_content_controls(self, doc: DocxDocument, section: str) -> list[FieldDef]:
        """Extract fields from DOCX structured document tags (SDTs)."""
        fields: list[FieldDef] = []
        body = doc.element.body
        sdt_index = 0
        for sdt in body.iter(qn('w:sdt')):
            alias_el = sdt.find(f'.//{qn("w:alias")}')
            tag_el = sdt.find(f'.//{qn("w:tag")}')
            placeholder_el = sdt.find(f'.//{qn("w:showingPlcHdr")}')

            alias = alias_el.get(qn('w:val')) if alias_el is not None else None
            tag = tag_el.get(qn('w:val')) if tag_el is not None else None

            label = alias or tag
            if not label:
                sdt_index += 1
                continue

            text_parts = [t.text for t in sdt.iter(qn('w:t')) if t.text]
            current_text = ' '.join(text_parts).strip()

            field_id = self._to_field_id(tag or label)
            fields.append(FieldDef(
                field_id=field_id,
                label=self._clean_label(label),
                field_type=self._infer_field_type(label),
                section=section,
                required=True,
                docx_location={"type": "content_control", "index": sdt_index, "tag": tag},
                placeholder_text=current_text,
            ))
            sdt_index += 1
        return fields

    def _get_fallback_cover_letter_fields(self) -> list[FieldDef]:
        """Hardcoded cover letter fields when automatic parsing finds nothing."""
        loc = lambda i: {"type": "paragraph", "index": i}
        return [
            FieldDef("project_title", "Project Title", "text", "Project Information", True, loc(0)),
            FieldDef("gs_id", "Gold Standard ID", "text", "Project Information", False, loc(1),
                      help_text="Leave blank if not yet assigned"),
            FieldDef("project_developer_name", "Project Developer Name", "text", "Project Developer", True, loc(2)),
            FieldDef("project_developer_address", "Address", "multiline", "Project Developer", True, loc(3)),
            FieldDef("project_developer_contact", "Contact Person", "text", "Project Developer", True, loc(4)),
            FieldDef("project_developer_email", "Email", "text", "Project Developer", True, loc(5)),
            FieldDef("project_developer_phone", "Phone", "text", "Project Developer", False, loc(6)),
            FieldDef("methodology", "Methodology / Protocol", "text", "Project Details", True, loc(7)),
            FieldDef("project_country", "Host Country", "text", "Project Details", True, loc(8)),
            FieldDef("project_scale", "Project Scale", "choice", "Project Details", True, loc(9),
                      help_text="Micro, Small, Regular, or Large"),
            FieldDef("crediting_period_start", "Crediting Period Start Date", "date", "Project Details", True, loc(10)),
            FieldDef("crediting_period_end", "Crediting Period End Date", "date", "Project Details", False, loc(11)),
            FieldDef("estimated_annual_credits", "Estimated Annual Emission Reductions (tCO2e)", "text",
                      "Project Details", False, loc(12)),
            FieldDef("project_description_summary", "Brief Project Description", "multiline",
                      "Project Summary", True, loc(13)),
            FieldDef("sustainability_contributions", "Sustainability Development Contributions", "multiline",
                      "Sustainability", False, loc(14)),
            FieldDef("stakeholder_engagement_summary", "Stakeholder Engagement Summary", "multiline",
                      "Stakeholder Engagement", False, loc(15)),
            FieldDef("documents_submitted", "Documents Submitted with this Cover Letter", "multiline",
                      "Submission Details", True, loc(16)),
            FieldDef("signatory_name", "Authorized Signatory Name", "text", "Declaration", True, loc(17)),
            FieldDef("signatory_title", "Title / Position", "text", "Declaration", True, loc(18)),
            FieldDef("signatory_organization", "Organization", "text", "Declaration", True, loc(19)),
            FieldDef("signature_date", "Date", "date", "Declaration", True, loc(20)),
            FieldDef("signature", "Signature", "signature", "Declaration", True, loc(21),
                      help_text="Leave blank — sign after export"),
        ]

    def _get_fallback_preliminary_review_fields(self) -> list[FieldDef]:
        """Bespoke fields for Preliminary Review when no template URL is available."""
        loc = lambda i: {"type": "paragraph", "index": i}
        return [
            FieldDef("project_title", "Project Title", "text", "Project Information", True, loc(0)),
            FieldDef("project_developer_name", "Project Developer / Proponent Name", "text", "Project Information", True, loc(1)),
            FieldDef("project_country", "Host Country", "text", "Project Information", True, loc(2)),
            FieldDef("methodology", "Methodology / Protocol", "text", "Project Information", True, loc(3)),
            FieldDef("project_scale", "Project Scale", "choice", "Project Information", True, loc(4),
                      help_text="Micro, Small, Regular, or Large"),
            FieldDef("project_summary", "Brief Project Summary", "multiline", "Project Description", True, loc(5)),
            FieldDef("expected_credits", "Expected Annual Emission Reductions (tCO2e)", "text", "Project Description", False, loc(6)),
            FieldDef("documents_list", "Documents Submitted with this Submission", "multiline", "Submission", True, loc(7)),
            FieldDef("signatory_name", "Authorized Signatory Name", "text", "Declaration", True, loc(8)),
            FieldDef("signatory_title", "Title / Position", "text", "Declaration", True, loc(9)),
            FieldDef("signature_date", "Date", "date", "Declaration", True, loc(10)),
        ]

    @staticmethod
    def get_section_contexts(template_type: str) -> dict[str, str]:
        """Return explanatory context text per section for questionnaire UI."""
        if template_type == TEMPLATE_TYPE_COVER_LETTER:
            return {
                "Project Information": "Basic project identification details required for all Gold Standard submissions.",
                "Project Developer": "Contact and organizational information for the project developer or implementing entity.",
                "Project Details": "Methodology, scale, crediting period, and estimated emission reductions.",
                "Project Summary": "A concise description of the project activities and expected outcomes.",
                "Sustainability": "How the project contributes to sustainable development goals (optional but recommended).",
                "Stakeholder Engagement": "Summary of consultations with local stakeholders (optional).",
                "Submission Details": "List of all documents being submitted with this cover letter.",
                "Declaration": "Authorized signatory information. Sign after export.",
            }
        if template_type == TEMPLATE_TYPE_PRELIMINARY_REVIEW:
            return {
                "Project Information": "Core project identification for the preliminary review screening.",
                "Project Description": "Brief overview of the project and expected emission reductions.",
                "Submission": "Documents accompanying this preliminary review submission.",
                "Declaration": "Authorized signatory details.",
            }
        return {}

    # ------------------------------------------------------------------
    # HTML preview generation
    # ------------------------------------------------------------------

    def _generate_html_preview(self, docx_bytes: bytes, field_schema: list[FieldDef]) -> str:
        """Convert DOCX to HTML via mammoth, then annotate field placeholders.

        Duplicate placeholder texts (e.g. "[Inert here]" across 5 entity tables)
        are annotated in document order, assigning the correct field_id to each
        successive occurrence.
        """
        import mammoth
        result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
        html = result.value

        placeholder_queues: dict[str, list[FieldDef]] = defaultdict(list)
        for f in field_schema:
            if f.placeholder_text:
                placeholder_queues[f.placeholder_text].append(f)

        for placeholder_text, field_list in placeholder_queues.items():
            escaped = re.escape(placeholder_text)
            counter = [0]

            def _make_replacer(p_text: str, f_list: list[FieldDef], ctr: list[int]):
                def replacer(_match):
                    f = f_list[min(ctr[0], len(f_list) - 1)]
                    ctr[0] += 1
                    return (
                        f'<span data-field-id="{f.field_id}" '
                        f'data-field-type="{f.field_type}" '
                        f'data-field-required="{str(f.required).lower()}" '
                        f'class="gs-field-placeholder">{p_text}</span>'
                    )
                return replacer

            html = re.sub(
                escaped,
                _make_replacer(placeholder_text, field_list, counter),
                html,
            )

        return html

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _to_field_id(text: str) -> str:
        cleaned = re.sub(r'[^a-zA-Z0-9\s]', '', text.lower())
        cleaned = re.sub(r'\s+', '_', cleaned.strip())
        return cleaned[:60] or "unknown_field"

    @staticmethod
    def _clean_label(text: str) -> str:
        text = re.sub(r'^(enter|insert|inert|click here to enter)\s+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'^the\s+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\s+(as|in|above|below)\s+.*$', '', text, flags=re.IGNORECASE)
        text = text.strip()
        if not text:
            return "Unknown"
        if text == text.upper() and len(text) > 1:
            return text.title()
        return text

    @staticmethod
    def _infer_field_type(text: str) -> str:
        lower = text.lower()
        if 'date' in lower:
            return 'date'
        if 'signature' in lower or 'sign' in lower:
            return 'signature'
        if any(kw in lower for kw in ('description', 'summary', 'address', 'notes', 'contributions', 'engagement')):
            return 'multiline'
        return 'text'

    async def _download_template(self, url: str) -> bytes:
        # 10s connect + 20s read — short enough to fail fast if GS site is slow
        timeout = httpx.Timeout(connect=10.0, read=20.0, write=5.0, pool=5.0)
        async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
            response = await client.get(url)
            response.raise_for_status()
            content = response.content
            # Sanity check: DOCX files start with PK (ZIP magic bytes)
            if not content.startswith(b'PK'):
                raise ValueError(
                    f"Downloaded file does not appear to be a valid DOCX "
                    f"(got {len(content)} bytes, starts with {content[:4]!r})"
                )
            return content

    async def _create_synthetic_template(self, template_type: str) -> GSTemplateVersion:
        """Create a placeholder template using the hardcoded fallback field schema.

        This is used when the GS website is unreachable or no URL is configured.
        Exports will use a generic blank document structure.
        """
        import io
        from docx import Document as DocxDocument

        if template_type == TEMPLATE_TYPE_PRELIMINARY_REVIEW:
            fallback_fields = self._get_fallback_preliminary_review_fields()
            doc_title = "Gold Standard Preliminary Review Submission"
        else:
            fallback_fields = self._get_fallback_cover_letter_fields()
            doc_title = "Gold Standard Cover Letter"
        field_schema_dicts = []
        for f in fallback_fields:
            d = asdict(f)
            if not d.get("placeholder_text"):
                d["placeholder_text"] = f"[{f.label}]"
            field_schema_dicts.append(d)

        # Build a minimal placeholder DOCX so fill_template has something to write into
        doc = DocxDocument()
        doc.add_heading(doc_title, 0)
        doc.add_paragraph("(Template loaded in offline mode — formatting may differ from the official GS template)")
        doc.add_paragraph("")
        for f in fallback_fields:
            doc.add_paragraph(f"{f.label}: {f.placeholder_text or f'[{f.label}]'}")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        docx_bytes = buf.read()

        file_hash = hashlib.sha256(docx_bytes).hexdigest()

        # Check if an identical synthetic template was already created
        existing = await self._get_version_by_hash(template_type, file_hash)
        if existing:
            return existing

        version_count = await self._count_versions(template_type)

        # Build a minimal HTML preview from the fallback fields
        html_lines = [f"<h1>{doc_title}</h1>",
                      "<p><em>(Offline mode — fill fields below)</em></p>"]
        for f in fallback_fields:
            placeholder = f.placeholder_text or f"[{f.label}]"
            html_lines.append(
                f'<p>{f.label}: <span data-field-id="{f.field_id}" '
                f'data-field-type="{f.field_type}" data-field-required="{str(f.required).lower()}" '
                f'class="gs-field-placeholder">{placeholder}</span></p>'
            )

        version = GSTemplateVersion(
            template_type=template_type,
            version_label=f"v{version_count + 1}-synthetic",
            source_url=TEMPLATE_URLS.get(template_type),
            file_hash=file_hash,
            file_bytes=docx_bytes,
            html_preview="\n".join(html_lines),
            field_schema=field_schema_dicts,
            status="draft",
        )
        self.db.add(version)
        await self.db.flush()
        await self.db.commit()
        await self.db.refresh(version)

        _update_check_cache.pop(template_type, None)
        logger.info("Created synthetic fallback template for %s (id=%s)", template_type, version.id)
        return version

    async def _get_version_by_hash(self, template_type: str, file_hash: str) -> Optional[GSTemplateVersion]:
        result = await self.db.execute(
            select(GSTemplateVersion).where(
                GSTemplateVersion.template_type == template_type,
                GSTemplateVersion.file_hash == file_hash,
            )
        )
        return result.scalar_one_or_none()

    async def _get_latest_by_status(self, template_type: str, status: str) -> Optional[GSTemplateVersion]:
        result = await self.db.execute(
            select(GSTemplateVersion).where(
                GSTemplateVersion.template_type == template_type,
                GSTemplateVersion.status == status,
            ).order_by(GSTemplateVersion.created_at.desc()).limit(1)
        )
        return result.scalar_one_or_none()

    async def _count_versions(self, template_type: str) -> int:
        from sqlalchemy import func
        result = await self.db.execute(
            select(func.count(GSTemplateVersion.id)).where(
                GSTemplateVersion.template_type == template_type,
            )
        )
        return result.scalar() or 0
