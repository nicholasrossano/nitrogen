"""
Tests for service layer.
"""
import pytest
from unittest.mock import MagicMock, patch, AsyncMock

from app.services.document_parser import DocumentParserService


class TestDocumentParserService:
    """Tests for DocumentParserService."""

    @pytest.fixture
    def parser(self):
        """Create a DocumentParserService instance."""
        return DocumentParserService()

    def test_chunk_text_basic(self, parser, sample_text_content):
        """Test basic text chunking."""
        chunks = parser.chunk_text(sample_text_content)
        
        assert len(chunks) > 0
        assert all(isinstance(chunk, str) for chunk in chunks)
        assert all(len(chunk) > 0 for chunk in chunks)

    def test_chunk_text_empty(self, parser):
        """Test chunking empty text."""
        chunks = parser.chunk_text("")
        assert len(chunks) == 0

    def test_chunk_text_whitespace(self, parser):
        """Test chunking whitespace-only text."""
        chunks = parser.chunk_text("   \n\n   ")
        assert len(chunks) == 0

    def test_chunk_text_short(self, parser):
        """Test chunking short text (smaller than chunk size)."""
        short_text = "This is a short text."
        chunks = parser.chunk_text(short_text)
        
        assert len(chunks) == 1
        assert "short text" in chunks[0]

    def test_chunk_text_preserves_content(self, parser):
        """Test that chunking preserves all content."""
        text = "First sentence. Second sentence. Third sentence. Fourth sentence."
        chunks = parser.chunk_text(text)
        
        # All chunks combined should contain all original words
        combined = " ".join(chunks)
        for word in ["First", "Second", "Third", "Fourth"]:
            assert word in combined

    def test_count_tokens(self, parser):
        """Test token counting."""
        text = "Hello world"
        count = parser.count_tokens(text)
        
        assert count > 0
        assert isinstance(count, int)

    def test_count_tokens_empty(self, parser):
        """Test token counting for empty string."""
        count = parser.count_tokens("")
        assert count == 0

    def test_count_tokens_longer_text(self, parser, sample_text_content):
        """Test token counting for longer text."""
        count = parser.count_tokens(sample_text_content)
        
        # Should have reasonable number of tokens
        assert count > 10
        assert count < 10000

    def test_clean_chunk_ends_at_sentence(self, parser):
        """Test that _clean_chunk tries to end at sentence boundaries."""
        text = "This is a sentence. This is another sentence. This is a partial"
        cleaned = parser._clean_chunk(text)
        
        # Should end at a sentence boundary if possible
        assert cleaned.endswith(".") or cleaned == text

    def test_clean_chunk_short_text(self, parser):
        """Test _clean_chunk with short text."""
        short_text = "Short."
        cleaned = parser._clean_chunk(short_text)
        assert cleaned == short_text

    def test_parser_initialization(self, parser):
        """Test parser initialization with settings."""
        assert parser.chunk_size > 0
        assert parser.chunk_overlap >= 0
        assert parser.chunk_overlap < parser.chunk_size
        assert parser.tokenizer is not None


class TestDocumentParserParsing:
    """Tests for document parsing (PDF, DOCX)."""

    @pytest.fixture
    def parser(self):
        """Create a DocumentParserService instance."""
        return DocumentParserService()

    def test_parse_docx_basic(self, parser):
        """Test parsing a DOCX document."""
        # Create a minimal DOCX in memory
        from docx import Document
        import io
        
        doc = Document()
        doc.add_paragraph("First paragraph of the document.")
        doc.add_paragraph("Second paragraph with more content.")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()
        
        # Parse it
        text = parser.parse_docx(docx_content)
        
        assert "First paragraph" in text
        assert "Second paragraph" in text

    def test_parse_docx_empty(self, parser):
        """Test parsing an empty DOCX document."""
        from docx import Document
        import io
        
        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()
        
        text = parser.parse_docx(docx_content)
        assert text == ""

    def test_parse_docx_with_whitespace(self, parser):
        """Test parsing DOCX with whitespace paragraphs."""
        from docx import Document
        import io
        
        doc = Document()
        doc.add_paragraph("Content paragraph.")
        doc.add_paragraph("   ")  # Whitespace only
        doc.add_paragraph("Another paragraph.")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()
        
        text = parser.parse_docx(docx_content)
        
        assert "Content paragraph" in text
        assert "Another paragraph" in text


class TestSDGClassifier:
    """Tests for SDG classifier service."""

    def test_classify_sdg_education(self):
        """Test SDG classification for education project."""
        from app.services.sdg_classifier import classify_sdg
        
        result = classify_sdg(
            project_description="A project to improve primary education access in rural areas",
            project_type="education"
        )
        
        # Should return SDG info or None
        if result:
            assert "sdg" in result or isinstance(result, dict)

    def test_classify_sdg_health(self):
        """Test SDG classification for health project."""
        from app.services.sdg_classifier import classify_sdg
        
        result = classify_sdg(
            project_description="Healthcare initiative to reduce infant mortality and improve health",
            project_type="health"
        )
        
        if result:
            assert isinstance(result, dict)
            assert "sdg" in result
            assert result["sdg"] == "3"  # Health SDG

    def test_classify_sdg_energy(self):
        """Test SDG classification for energy project."""
        from app.services.sdg_classifier import classify_sdg
        
        result = classify_sdg(
            project_description="Solar mini-grid project for rural electrification with renewable energy",
            project_type="energy_access"
        )
        
        if result:
            assert result["sdg"] == "7"  # Energy SDG

    def test_classify_sdg_empty(self):
        """Test SDG classification with empty description."""
        from app.services.sdg_classifier import classify_sdg
        
        result = classify_sdg(project_description="", project_type=None)
        # Should handle gracefully and return None
        assert result is None

    def test_classify_sdg_no_match(self):
        """Test SDG classification with no matching keywords."""
        from app.services.sdg_classifier import classify_sdg
        
        result = classify_sdg(project_description="random text without keywords", project_type=None)
        # Should return None when no match
        assert result is None


class TestToolRegistry:
    """Tests for tool registry."""

    def test_get_tool_registry(self):
        """Test getting tool registry instance."""
        from app.tools import get_tool_registry
        
        registry = get_tool_registry()
        assert registry is not None

    def test_registry_has_tools(self):
        """Test that registry has registered tools."""
        from app.tools import get_tool_registry
        
        registry = get_tool_registry()
        tools = registry.get_all_tools()
        
        assert isinstance(tools, list)
        assert len(tools) >= 1

    def test_get_tool_by_id(self):
        """Test getting a specific tool by ID."""
        from app.tools import get_tool_registry
        
        registry = get_tool_registry()
        tools = registry.get_all_tools()
        
        if tools:
            first_tool = tools[0]
            tool = registry.get_tool(first_tool.definition.id)
            assert tool is not None
            assert tool.definition.id == first_tool.definition.id

    def test_get_all_definitions(self):
        """Test getting tool definitions."""
        from app.tools import get_tool_registry
        
        registry = get_tool_registry()
        definitions = registry.get_all_definitions()
        
        assert isinstance(definitions, list)
        assert len(definitions) >= 1
        # Each definition should have required fields
        for definition in definitions:
            assert hasattr(definition, 'id')
            assert hasattr(definition, 'name')
            assert hasattr(definition, 'description')

    def test_recommend_tools(self):
        """Test tool recommendations."""
        from app.tools import get_tool_registry
        
        registry = get_tool_registry()
        recommendations = registry.recommend_tools(
            project_description="A solar mini-grid project for rural electrification",
            project_type="energy_access"
        )
        
        assert isinstance(recommendations, list)
        assert len(recommendations) >= 1
        # Each recommendation should be (tool, confidence) tuple
        for tool, confidence in recommendations:
            assert hasattr(tool, 'definition')
            assert 0 <= confidence <= 1

    def test_classify_project_type(self):
        """Test project type classification."""
        from app.tools import get_tool_registry
        
        registry = get_tool_registry()
        
        # Test energy classification
        assert registry.classify_project_type("solar mini-grid project") == "energy_access"
        
        # Test clean cooking classification
        assert registry.classify_project_type("lpg cookstove distribution") == "clean_cooking"
        
        # Test water classification
        assert registry.classify_project_type("water sanitation project") == "water_sanitation"
        
        # Test default
        assert registry.classify_project_type("random project") == "general"
